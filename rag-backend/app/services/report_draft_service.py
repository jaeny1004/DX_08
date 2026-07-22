from __future__ import annotations

import csv
import json
import re
import shutil
import subprocess
import uuid
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill

BACKEND_ROOT = Path(__file__).resolve().parents[2]
DATA_ROOT = BACKEND_ROOT / "data"
CANDIDATE_GEOJSON = DATA_ROOT / "final_ui_candidate_v4.geojson"
HISTORICAL_ROOT = DATA_ROOT / "generated_reports"
DRAFT_ROOT = DATA_ROOT / "generated_drafts"

REPORT_LABELS = {
    "prediction": "신규 확산위험 분석 보고서",
    "field_survey_plan": "현장 예찰 계획서",
    "field_survey_result": "현장 예찰 결과보고서",
    "control_plan": "방제 검토 계획서",
    "integrated": "예측·예찰·방제 통합 보고서",
}

HISTORY_DIR = {
    "prediction": "prediction_30",
    "field_survey_plan": "field_survey_30",
    "field_survey_result": "field_survey_30",
    "control_plan": "control_30",
    "integrated": "prediction_30",
}

ALIASES = {
    "grid_id": ["center_grid_id", "grid_id", "id", "GRID_ID"],
    "sido": ["sido_name", "sido", "ctpv_nm", "SIDO_NM"],
    "sigungu": ["sigungu_name", "sigungu", "sgg_nm", "SIGUNGU_NM"],
    "risk_score": ["risk_score", "riskScore", "pred_score", "prediction_score", "probability"],
    "risk_grade": ["risk_grade", "riskGrade", "risk_level", "risk_label"],
    "priority_score": ["priority_score", "priorityScore", "survey_priority_score"],
    "priority_grade": ["priority_grade", "priorityGrade", "priority_label", "priority_stage_label"],
    "pine_ratio": ["pine_ratio", "pineRatio"],
    "infection_pressure": ["infection_pressure", "recent_infection_pressure", "infect_pressure"],
    "access_score": ["access_score_v3", "access_score", "accessibility_score"],
}


def _pick(props: dict[str, Any], key: str) -> Any:
    for name in ALIASES[key]:
        value = props.get(name)
        if value is not None and str(value).strip() != "":
            return value
    return None


def _text(value: Any) -> str:
    return "" if value is None else str(value).strip()


def _float(value: Any) -> float | None:
    try:
        return float(str(value).replace(",", "").strip())
    except (TypeError, ValueError):
        return None


def _mean(values: list[float]) -> float | None:
    return round(sum(values) / len(values), 2) if values else None


def _max(values: list[float]) -> float | None:
    return round(max(values), 2) if values else None


def _normalize_grade(value: Any) -> str:
    text = _text(value)
    return {"매우높음": "매우 높음", "매우낮음": "매우 낮음"}.get(text, text)


def _safe_name(value: str) -> str:
    return re.sub(r'[\\/:*?"<>|]+', "_", value).strip()[:120] or "report"


def _load_features() -> list[dict[str, Any]]:
    if not CANDIDATE_GEOJSON.is_file():
        raise FileNotFoundError(f"현재 분석 데이터 파일이 없습니다: {CANDIDATE_GEOJSON}")
    with CANDIDATE_GEOJSON.open("r", encoding="utf-8") as file:
        payload = json.load(file)
    features = payload.get("features")
    if not isinstance(features, list):
        raise ValueError("GeoJSON features 배열을 확인할 수 없습니다.")
    return features


def collect_current_data(sido_name: str, sigungu_name: str, grid_ids: list[str]) -> dict[str, Any]:
    selected: list[dict[str, Any]] = []
    wanted = set(grid_ids)

    for feature in _load_features():
        props = feature.get("properties") or {}
        if not isinstance(props, dict):
            continue
        if sido_name and _text(_pick(props, "sido")) != sido_name:
            continue
        if sigungu_name and _text(_pick(props, "sigungu")) != sigungu_name:
            continue
        if wanted and _text(_pick(props, "grid_id")) not in wanted:
            continue
        selected.append(props)

    if not selected:
        raise ValueError("선택 조건에 해당하는 현재 분석 데이터가 없습니다.")

    risks = [v for v in (_float(_pick(p, "risk_score")) for p in selected) if v is not None]
    priorities = [v for v in (_float(_pick(p, "priority_score")) for p in selected) if v is not None]
    pine = [v for v in (_float(_pick(p, "pine_ratio")) for p in selected) if v is not None]
    pressure = [v for v in (_float(_pick(p, "infection_pressure")) for p in selected) if v is not None]
    access = [v for v in (_float(_pick(p, "access_score")) for p in selected) if v is not None]

    risk_dist: dict[str, int] = {}
    priority_dist: dict[str, int] = {}
    ids: list[str] = []
    for props in selected:
        gid = _text(_pick(props, "grid_id"))
        if gid:
            ids.append(gid)
        rg = _normalize_grade(_pick(props, "risk_grade"))
        pg = _text(_pick(props, "priority_grade"))
        if rg:
            risk_dist[rg] = risk_dist.get(rg, 0) + 1
        if pg:
            priority_dist[pg] = priority_dist.get(pg, 0) + 1

    return {
        "selected_grid_count": len(selected),
        "grid_ids": ids,
        "average_risk_score": _mean(risks),
        "maximum_risk_score": _max(risks),
        "average_priority_score": _mean(priorities),
        "maximum_priority_score": _max(priorities),
        "average_pine_ratio": _mean(pine),
        "average_infection_pressure": _mean(pressure),
        "average_access_score": _mean(access),
        "risk_grade_distribution": risk_dist,
        "priority_grade_distribution": priority_dist,
    }


def find_references(report_type: str, sido: str, sigungu: str, document_nos: list[str]) -> list[dict[str, Any]]:
    path = HISTORICAL_ROOT / HISTORY_DIR[report_type] / "문서목록.csv"
    if not path.is_file():
        return []
    with path.open("r", encoding="utf-8-sig", newline="") as file:
        rows = [{k: (v or "").strip() for k, v in row.items()} for row in csv.DictReader(file)]

    if document_nos:
        chosen = [r for r in rows if r.get("document_no") in set(document_nos)]
    else:
        chosen = [r for r in rows if r.get("sido_name") == sido and r.get("sigungu_name") == sigungu]
        chosen = chosen or [r for r in rows if r.get("sido_name") == sido] or rows

    return [
        {
            "document_no": r.get("document_no", ""),
            "file_name": r.get("file_name", ""),
            "year": r.get("year", ""),
            "sido_name": r.get("sido_name", ""),
            "sigungu_name": r.get("sigungu_name", ""),
            "center_grid_id": r.get("center_grid_id", ""),
            "risk_grade": _normalize_grade(r.get("risk_grade")),
            "priority_grade": r.get("priority_grade", ""),
        }
        for r in chosen[:3]
    ]


def _distribution_text(values: dict[str, int]) -> str:
    if not values:
        return "등급 정보 없음"
    return ", ".join(f"{k} {v}개" for k, v in sorted(values.items(), key=lambda x: x[1], reverse=True))


def build_sections(payload: dict[str, Any], summary: dict[str, Any], refs: list[dict[str, Any]]) -> list[dict[str, str]]:
    region = f"{payload['sido_name']} {payload.get('sigungu_name', '')}".strip()
    ref_text = ", ".join(f"문서번호 {r['document_no']}({r['year']}년 {r['sido_name']} {r['sigungu_name']})" for r in refs) or "자동 참고문서 없음"
    include = payload.get("include_sections", {})

    sections = [
        {"key": "overview", "heading": "1. 분석 개요", "content": f"본 문서는 {payload['start_date']}부터 {payload['end_date']}까지 {region}을 대상으로 500m 격자 기반 신규 확산위험 후보와 예찰 우선순위를 검토하기 위한 행정 초안이다. 총 {summary['selected_grid_count']}개 격자를 분석하였다."},
        {"key": "basis", "heading": "2. 활용 데이터 및 작성 기준", "content": f"수치는 현재 플랫폼 분석 결과에서 직접 집계했으며 문서 구성과 행정 표현은 과거 보고서를 참고하였다. 참고 문서: {ref_text}. 검경 전 감염 확정 표현은 사용하지 않는다."},
    ]

    if include.get("risk_summary", True):
        sections.append({"key": "risk", "heading": "3. 신규 확산위험 후보 분석", "content": f"평균 위험도 점수는 {summary.get('average_risk_score') or '-'}점, 최대 위험도 점수는 {summary.get('maximum_risk_score') or '-'}점이다. 등급 분포는 {_distribution_text(summary.get('risk_grade_distribution', {}))}이다. 상위 위험 격자는 우선 예찰 검토지역으로 관리한다."})
    if include.get("priority_summary", True):
        sections.append({"key": "priority", "heading": "4. 예찰 우선순위 검토", "content": f"평균 예찰 우선순위 점수는 {summary.get('average_priority_score') or '-'}점, 최대 점수는 {summary.get('maximum_priority_score') or '-'}점이다. 분포는 {_distribution_text(summary.get('priority_grade_distribution', {}))}이다."})
    if include.get("infection_history", True):
        sections.append({"key": "history", "heading": "5. 감염 발생 이력 참고", "content": "과거 발생 이력은 현재 감염 확정 근거가 아니라 확산압력과 현장 확인 우선순위를 판단하기 위한 참고자료로 활용한다."})
    if include.get("workforce_plan", False):
        sections.append({"key": "workforce", "heading": "6. 현장 예찰 운영 검토", "content": "위험도와 예찰 우선순위가 높은 격자부터 배정하고, 접근성이 낮거나 범위가 넓은 지역은 드론 예찰을 선행한다."})
    if include.get("control_scenario", False):
        sections.append({"key": "control", "heading": "7. 방제 검토 방향", "content": "방제방법은 현장 확인과 검경 결과 이후 결정하며, 현재 단계에서는 파쇄·훈증·예방나무주사 등을 비교하는 방제 검토 계획으로 작성한다."})

    note = payload.get("user_notes", "").strip()
    sections.append({"key": "actions", "heading": "8. 행정 조치 제안", "content": "가. 상위 위험 격자 우선 예찰 일정 수립\n나. 현장 결과와 시료 채취 내역 표준화 입력\n다. 예측·현장예찰·방제 문서 연계\n라. 검경 전 감염 확정 표현 사용 금지" + (f"\n마. 담당자 메모: {note}" if note else "")})
    return sections


def _draft_path(draft_id: str) -> Path:
    return DRAFT_ROOT / draft_id / "draft.json"


def create_draft(payload: dict[str, Any], created_by: str) -> dict[str, Any]:
    DRAFT_ROOT.mkdir(parents=True, exist_ok=True)
    now = datetime.now().astimezone().isoformat(timespec="seconds")
    draft_id = f"DRAFT-{datetime.now():%Y%m%d}-{uuid.uuid4().hex[:8].upper()}"
    grid_ids = [str(v).strip() for v in payload.get("center_grid_ids", []) if str(v).strip()]
    ref_nos = [str(v).strip() for v in payload.get("reference_document_nos", []) if str(v).strip()]
    summary = collect_current_data(payload["sido_name"], payload.get("sigungu_name", ""), grid_ids)
    refs = find_references(payload["report_type"], payload["sido_name"], payload.get("sigungu_name", ""), ref_nos)
    title = payload.get("title", "").strip() or f"{payload['year']}년 {payload.get('sigungu_name') or payload['sido_name']} {REPORT_LABELS[payload['report_type']]}"

    draft = {
        "draft_id": draft_id,
        "report_type": payload["report_type"],
        "title": title,
        "status": "draft",
        "created_at": now,
        "updated_at": now,
        "created_by": created_by,
        "year": payload["year"],
        "start_date": payload["start_date"],
        "end_date": payload["end_date"],
        "sido_name": payload["sido_name"],
        "sigungu_name": payload.get("sigungu_name", ""),
        "center_grid_ids": grid_ids,
        "reference_document_nos": ref_nos,
        "reference_reports": refs,
        "data_summary": summary,
        "sections": build_sections(payload, summary, refs),
        "user_notes": payload.get("user_notes", "").strip(),
    }
    save_draft(draft)
    return draft


def save_draft(draft: dict[str, Any]) -> None:
    path = _draft_path(draft["draft_id"])
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", encoding="utf-8") as file:
        json.dump(draft, file, ensure_ascii=False, indent=2)


def load_draft(draft_id: str) -> dict[str, Any]:
    path = _draft_path(draft_id)
    if not path.is_file():
        raise FileNotFoundError(f"초안을 찾을 수 없습니다: {draft_id}")
    with path.open("r", encoding="utf-8") as file:
        return json.load(file)


def update_draft(draft_id: str, payload: dict[str, Any]) -> dict[str, Any]:
    draft = load_draft(draft_id)
    if payload.get("title") is not None:
        draft["title"] = payload["title"].strip() or draft["title"]
    if payload.get("status") is not None:
        draft["status"] = payload["status"]
    if payload.get("sections") is not None:
        draft["sections"] = payload["sections"]
    draft["updated_at"] = datetime.now().astimezone().isoformat(timespec="seconds")
    save_draft(draft)
    return draft


def build_docx(draft: dict[str, Any]) -> Path:
    directory = _draft_path(draft["draft_id"]).parent
    path = directory / f"{_safe_name(draft['title'])}.docx"
    document = Document()
    document.styles["Normal"].font.name = "Malgun Gothic"
    document.styles["Normal"].font.size = Pt(10)
    title = document.add_heading(draft["title"], level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph(f"문서번호: {draft['draft_id']}\n대상지역: {draft['sido_name']} {draft['sigungu_name']}\n대상기간: {draft['start_date']} ~ {draft['end_date']}\n문서상태: 행정 검토용 초안")
    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "핵심 지표"
    table.rows[0].cells[1].text = "값"
    for key in ("selected_grid_count", "average_risk_score", "maximum_risk_score", "average_priority_score", "maximum_priority_score", "average_pine_ratio", "average_infection_pressure", "average_access_score"):
        cells = table.add_row().cells
        cells[0].text = key
        cells[1].text = str(draft["data_summary"].get(key) or "-")
    for section in draft["sections"]:
        document.add_heading(section["heading"], level=1)
        for line in section["content"].splitlines():
            document.add_paragraph(line)
    document.add_paragraph("※ 본 문서는 AI 기반 분석 결과와 과거 행정 양식을 활용한 초안이며 최종 결재 전 담당자 검토가 필요합니다.")
    document.save(path)
    return path


def build_xlsx(draft: dict[str, Any]) -> Path:
    directory = _draft_path(draft["draft_id"]).parent
    path = directory / f"{_safe_name(draft['title'])}.xlsx"
    wb = Workbook()
    ws = wb.active
    ws.title = "보고서 요약"
    ws.append(["항목", "값"])
    ws["A1"].font = ws["B1"].font = Font(bold=True)
    ws["A1"].fill = ws["B1"].fill = PatternFill("solid", fgColor="D9EAD3")
    for key, value in [("문서번호", draft["draft_id"]), ("제목", draft["title"]), ("대상지역", f"{draft['sido_name']} {draft['sigungu_name']}"), ("대상기간", f"{draft['start_date']} ~ {draft['end_date']}"), ("상태", draft["status"])]:
        ws.append([key, value])
    stats = wb.create_sheet("분석 지표")
    stats.append(["항목", "값"])
    for key, value in draft["data_summary"].items():
        stats.append([key, json.dumps(value, ensure_ascii=False) if isinstance(value, (dict, list)) else value])
    body = wb.create_sheet("문서 본문")
    body.append(["구분", "제목", "내용"])
    for section in draft["sections"]:
        body.append([section["key"], section["heading"], section["content"]])
    refs = wb.create_sheet("참고 과거문서")
    refs.append(["문서번호", "연도", "시도", "시군구", "중심 격자", "위험도", "예찰 우선순위", "파일명"])
    for row in draft["reference_reports"]:
        refs.append([row.get("document_no"), row.get("year"), row.get("sido_name"), row.get("sigungu_name"), row.get("center_grid_id"), row.get("risk_grade"), row.get("priority_grade"), row.get("file_name")])
    for sheet in wb.worksheets:
        for row in sheet.iter_rows():
            for cell in row:
                cell.alignment = Alignment(vertical="top", wrap_text=True)
        sheet.column_dimensions["A"].width = 24
        sheet.column_dimensions["B"].width = 50
        sheet.column_dimensions["C"].width = 80
    wb.save(path)
    return path


def build_pdf(draft: dict[str, Any]) -> Path:
    docx = build_docx(draft)
    libreoffice = shutil.which("libreoffice") or shutil.which("soffice")
    if not libreoffice:
        raise RuntimeError("LibreOffice가 없어 PDF 변환을 할 수 없습니다. DOCX는 정상 생성됐습니다.")
    result = subprocess.run([libreoffice, "--headless", "--convert-to", "pdf", "--outdir", str(docx.parent), str(docx)], capture_output=True, text=True, timeout=120, check=False)
    pdf = docx.with_suffix(".pdf")
    if result.returncode != 0 or not pdf.is_file():
        raise RuntimeError(f"PDF 변환 실패: {result.stderr.strip()}")
    return pdf
