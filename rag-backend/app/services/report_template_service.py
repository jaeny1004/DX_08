from __future__ import annotations

import csv
import json
import math
import re
import shutil
import subprocess
import tempfile
import zipfile
from datetime import date, datetime, timedelta
from pathlib import Path
from typing import Any

from docx import Document
from PIL import Image, ImageDraw, ImageFont

from app.services.report_draft_service import (
    ALIASES,
    CANDIDATE_GEOJSON,
    DATA_ROOT,
    REPORT_LABELS,
    _pick,
    _text,
    load_draft,
    save_draft,
)

TEMPLATE_ROOT = DATA_ROOT / "report_templates"
GENERATED_REPORT_ROOT = DATA_ROOT / "generated_reports"

TEMPLATE_FILES = {
    "prediction": "[양식]소나무재선충병 발생 예측 보고서_빈양식.docx",
    "field_survey": "[양식]소나무재선충병 현장 예찰 보고서_빈양식.docx",
    "control": "[양식]소나무재선충병 방제 보고서_빈양식.docx",
}

REPORT_DIRECTORIES = {
    "prediction": "prediction_30",
    "field_survey": "field_survey_30",
    "control": "control_30",
}


def _value(value: Any, suffix: str = "") -> str:
    if value is None or value == "":
        return "-"
    if isinstance(value, float):
        return f"{value:.2f}".rstrip("0").rstrip(".") + suffix
    return str(value) + suffix


def _percent(value: Any) -> str:
    if not isinstance(value, (int, float)):
        return "-"
    number = float(value) * 100 if float(value) <= 1 else float(value)
    return _value(round(number, 1))


def _days(start: str, end: str) -> int:
    try:
        return (date.fromisoformat(end) - date.fromisoformat(start)).days + 1
    except ValueError:
        return 1


def _replace_paragraph(paragraph, replacements: dict[str, str], exact: dict[str, str]) -> None:
    original = paragraph.text
    updated = exact.get(original, original)
    for source, target in sorted(replacements.items(), key=lambda item: len(item[0]), reverse=True):
        updated = updated.replace(source, target)
    if updated == original:
        return
    if paragraph.runs:
        paragraph.runs[0].text = updated
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.text = updated


def _all_paragraphs(document: Document):
    for paragraph in document.paragraphs:
        yield paragraph
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph
    for section in document.sections:
        for paragraph in section.header.paragraphs:
            yield paragraph
        for paragraph in section.footer.paragraphs:
            yield paragraph


def _common_replacements(draft: dict[str, Any]) -> dict[str, str]:
    center = draft["data_summary"].get("center_grid", {})
    region = f"{draft['sido_name']} {draft['sigungu_name']}".strip()
    period = f"{draft['start_date']} ~ {draft['end_date']}"
    risk = _value(center.get("risk_score"))
    risk_grade = _value(center.get("risk_grade"))
    return {
        "[작성일]": datetime.now().strftime("%Y-%m-%d"),
        "[시작일]": draft["start_date"],
        "[종료일]": draft["end_date"],
        "[일수]": str(_days(draft["start_date"], draft["end_date"])),
        "[지역]": region,
        "[기간]": period,
        "[주소]": region,
        "[격자 ID]": str(center.get("grid_id", "-")),
        "[점수]": risk,
        "[등급]": risk_grade,
        "[비율]": _percent(center.get("pine_ratio")),
        "[거리]": _value(center.get("road_distance")),
        "[도로 유형]": _value(center.get("road_type")),
        "[해당 여부]": _value(center.get("environment_flag")),
        "[위도]": _value(center.get("latitude")),
        "[경도]": _value(center.get("longitude")),
        "[일시]": draft["end_date"],
        "-지역, 기간-": f"-{region}, {period}-",
        "[소속·조]": "관할 산림 담당 부서",
        "[성명]": draft.get("created_by", "담당자"),
        "[정보]": f"중심 격자 {center.get('grid_id', '-')}",
        "[날씨]": "현장 조사 시 입력",
        "[수종]": "소나무류",
        "[면적]": "현장 확인 후 산정",
        "[포함 범위]": region,
        "[단계 수]": "5",
        "[단계]": "위험도",
        "[일자]": draft["end_date"],
        "[대상 격자]": str(center.get("grid_id", "-")),
    }


def _prediction_exact(draft: dict[str, Any]) -> dict[str, str]:
    center = draft["data_summary"].get("center_grid", {})
    neighbors = draft["data_summary"].get("neighbor_grids", [])
    region = f"{draft['sido_name']} {draft['sigungu_name']}".strip()
    risk = _value(center.get("risk_score"))
    risk_grade = _value(center.get("risk_grade"))
    priority = _value(center.get("priority_score"))
    priority_grade = _value(center.get("priority_grade"))
    pressure = _value(center.get("infection_pressure"))
    access = _value(center.get("access_score"))
    road_distance = _value(center.get("road_distance"))
    road_type = _value(center.get("road_type"))
    action = center.get("recommended_action") or f"{priority_grade} 기준으로 현장 확인 순서를 검토"
    neighbor_text = ", ".join(str(n.get("grid_id")) for n in neighbors if n.get("grid_id")) or "인접 후보격자 없음"
    exact = {
        "❍ (분석 배경) [입력]": f"❍ (분석 배경) {region}의 500m 격자별 산림·지형·기후·접근성 및 과거 감염 발생 이력을 결합하여 신규 확산위험 후보를 분석함",
        "❍ (분석 목적) [입력]": "❍ (분석 목적) 감염 확정이 아닌 신규 발생 후보지역을 사전에 선별하고 현장 예찰 자원을 우선 배치하기 위함",
        "❍ (활용 목적) [입력]": "❍ (활용 목적) 위험도와 예찰 우선순위를 함께 검토하여 현장 확인·드론 예찰·행정 보고 순서를 지원함",
        "❍ (위험 점수) 종합 위험도 스코어 [점수]점 (전체 [단계 수]단계 중 [단계]단계 ‘[등급]’ 수준)": f"❍ (위험 점수) 신규 확산위험 점수 {risk}점 (5단계 중 ‘{risk_grade}’ 수준)",
        "― 최근 감염압력 ‘[점수]’점": f"― 최근 감염압력 ‘{pressure}’점",
        "❍ 예찰 우선순위 ‘[점수]’점([등급])": f"❍ 예찰 우선순위 ‘{priority}’점({priority_grade})",
        "❍ 접근성 ‘[점수]’점": f"❍ 접근성 ‘{access}’점",
        "― 도로까지의 거리 ‘[거리]’m([도로 유형])": f"― 도로까지의 거리 ‘{road_distance}’m({road_type})",
        "❍ 향후 [기간] 내 [방향·대상 지역]으로의 신규 확산위험 [결과] 예측": f"❍ 중심 격자 주변 {neighbor_text}를 인접 우선 검토 대상으로 설정하고 위험도 변화를 지속 관찰",
        "❍ 기준 격자 ID [격자 ID] / 방향 [방향]": f"❍ 기준 격자 ID {center.get('grid_id', '-')} / 인접 후보격자 4개 기준",
        "❍ (시뮬레이션 종합 의견) [입력]": f"❍ (종합 의견) 중심 격자는 위험도 {risk}점({risk_grade}), 예찰 우선순위 {priority}점({priority_grade})으로 확인되며, 인접 후보격자와 함께 현장 확인이 필요함",
        "❍ 1단계: [대응 단계명]": "❍ 1단계: 우선 예찰 검토지역 지정",
        "― (실행 조건) [입력]": f"― (실행 조건) 위험도 {risk_grade} 및 예찰 우선순위 {priority_grade} 기준",
        "― (실행 계획) [일자] / [대상 격자] / [실행 내용]": f"― (실행 계획) {draft['end_date']} / {center.get('grid_id', '-')} / 현장 확인 및 사진·좌표 기록",
        "❍ 2단계: [대응 단계명]": "❍ 2단계: 인접 후보격자 연계 예찰",
        "― (대상 조직·알림 내용) [입력]": f"― (대상 조직·알림 내용) {region} 담당자에게 중심 격자와 인접 후보격자 우선 확인 알림",
        "― (현장 확인 계획) [입력]": f"― (현장 확인 계획) 접근성 {access}점과 도로거리 {road_distance}m를 고려하여 현장 또는 드론 예찰 배정",
        "❍ 3단계: [대응 단계명]": "❍ 3단계: 확인 결과 기반 후속 조치",
        "― (방제·수종전환 검토 사항) [입력]": f"― (방제·수종전환 검토 사항) 현장 확인 및 검경 결과에 따라 방제 필요성과 수종전환 가능성을 별도 검토",
        "― (행정 연계 및 후속 계획) [입력]": f"― (행정 연계 및 후속 계획) 확인 결과를 플랫폼 이력과 보고서에 등록하고 후속 예찰 일정을 관리 ({action})",
    }
    for index in range(4):
        token = f"― 격자 ID [인접 격자 {index + 1}]"
        if index < len(neighbors):
            n = neighbors[index]
            exact[token] = (
                f"― 격자 ID {n.get('grid_id', '-')} / {n.get('direction', '-')} / "
                f"약 {_value(n.get('distance_km'))}km / 위험도 {_value(n.get('risk_score'))}점"
                f"({n.get('risk_grade') or '-'}) / 예찰 {n.get('priority_grade') or '-'}"
            )
        else:
            exact[token] = "― 추가 인접 후보격자 없음"
    return exact


def _field_exact(draft: dict[str, Any]) -> dict[str, str]:
    center = draft["data_summary"].get("center_grid", {})
    risk = _value(center.get("risk_score"))
    grade = _value(center.get("risk_grade"))
    priority = _value(center.get("priority_score"))
    pgrade = _value(center.get("priority_grade"))
    return {
        "❍ (발견 경로) [입력]": "❍ (발견 경로) AI 신규 확산위험 후보 및 예찰 우선순위 지도에서 중심 격자 선택",
        "❍ (AI 분석 결과) [점수·등급 또는 판단 내용]": f"❍ (AI 분석 결과) 신규 확산위험 {risk}점({grade}), 예찰 우선순위 {priority}점({pgrade})",
        "❍ (종합 판단) [현장 확인 필요 여부 및 판단 근거 입력]": "❍ (종합 판단) 감염 확정값이 아니므로 현장 확인 및 필요 시 시료 검경이 필요함",
        "❍ (수종/수량) [수종] / 총 [수량]본 ([세부 분류])": "❍ (수종/수량) 소나무류 / 현장 조사 후 수량 및 피해 단계 입력",
        "❍ (변색 단계) [관찰 내용 및 단계]": "❍ (변색 단계) 현장 관찰 후 정상·초기 변색·진행 변색·고사 단계 기록",
        "❍ (매개충 흔적) [관찰 내용]": "❍ (매개충 흔적) 탈출공·후식 흔적 등 현장 관찰 결과 입력",
        "❍ (수피·목질부 관찰) [관찰 내용]": "❍ (수피·목질부 관찰) 수피 상태와 목질부 변색 여부 기록",
        "❍ “[현장 조사자 의견 입력]”": f"❍ “중심 격자 {center.get('grid_id', '-')}는 현장 확인 필요 대상이며 최종 판단은 조사·검경 결과를 따름”",
        "❍ (시료 채취) [시료 종류·수량] / QR코드 [번호]": "❍ (시료 채취) 현장 필요 시 목편·가지 시료 수량과 QR코드 입력",
        "❍ (시스템 연동) [처리 결과 입력]": "❍ (시스템 연동) 위치·사진·조사자 의견·시료 정보를 현장 이력에 등록",
        "❍ [일자]: [검경 의뢰 기관 및 진행 상태]": f"❍ {draft['end_date']}: 관할 검경기관 의뢰 여부 및 진행 상태 입력",
        "❍ (후속 조치 계획) [현장 확인·검경 결과에 따른 조치 입력]": "❍ (후속 조치 계획) 확인 결과에 따라 재예찰·검경·방제 검토 단계로 연계",
    }


def _control_exact(draft: dict[str, Any]) -> dict[str, str]:
    center = draft["data_summary"].get("center_grid", {})
    risk = _value(center.get("risk_score"))
    grade = _value(center.get("risk_grade"))
    return {
        "❍ (방 제 자) [소속·조] (단원: [성명])": f"❍ (방 제 자) 관할 산림 담당 부서 (담당자: {draft.get('created_by', '담당자')})",
        "❍ (방제 면적) 총 [면적] ha ([포함 범위])": f"❍ (방제 면적) 현장 확정 후 입력 ({draft['sido_name']} {draft['sigungu_name']} 중심 격자 기준)",
        "❍ (대상 수량) 최종 검경 확진목 [수량]본 및 감염 우려 피해목 [수량]본 (총 [수량]본)": "❍ (대상 수량) 현장 확인 및 검경 확정 후 확진목·피해목 수량 입력",
        "❍ (파쇄 처리) [수량]본 / [처리 방법 및 규격]": "❍ (파쇄 처리) 적용 시 처리 수량·파쇄 규격·완료 일시 입력",
        "❍ (훈증 처리) [수량]본 / [처리 사유 및 방법]": "❍ (훈증 처리) 적용 시 처리 수량·사유·피복 및 약제 정보 입력",
        "― 타포린 피복 일련번호: [번호]": "― 타포린 피복 일련번호: 훈증 적용 시 입력",
        "❍ (작업 면적) [범위 및 대상] (약 [면적]ha)": "❍ (작업 면적) 예방나무주사 적용 범위와 대상 면적 현장 확정 후 입력",
        "❍ (주입 실적) [수종] 약 [수량]본 / [약제 및 처리 내용]": "❍ (주입 실적) 수종·본수·약제명·주입량 및 작업 결과 입력",
        "❍ (천공 규격) [직경]mm, 깊이 [깊이]cm / [작업 방법]": "❍ (천공 규격) 사용 약제 지침과 현장 작업 기준에 따라 입력",
        "❍ (방제 전) [현장 상태 입력]": f"❍ (방제 전) 중심 격자 {center.get('grid_id', '-')} / 신규 확산위험 {risk}점({grade}) / 현장 상태 확인 필요",
        "❍ (방제 후) [조치 결과 입력]": "❍ (방제 후) 처리 수량·방법·잔재물 상태·완료 사진을 등록",
        "― 위험도 변화: 기존 [점수]점([등급]) → 방제 완료 후 [점수]점([등급])": f"― 위험도 변화: 기존 {risk}점({grade}) → 방제 완료 후 후속 예측 결과로 갱신",
        "― 훈증 더미 [개소] 위치 좌표 등록 및 [주기] 모니터링 계획": "― 훈증 더미 적용 시 개소·위치 좌표를 등록하고 정기 모니터링 계획 수립",
        "― (잔재물 관리 조치) [입력]": "― (잔재물 관리 조치) 파쇄물·훈증더미·운반 잔재물 처리 상태 기록",
        "❍ (보고 및 결재) [보고 대상·승인 절차 입력]": "❍ (보고 및 결재) 담당 부서 검토 후 내부 결재 및 보고서 등록",
        "― 1차 모니터링 예정일: [일자] ([방법])": f"― 1차 모니터링 예정일: {(date.fromisoformat(draft['end_date']) + timedelta(days=30)).isoformat()} (현장 재확인 또는 드론 예찰)",
        "― (후속 사업 및 행정 연계 계획) [입력]": "― (후속 사업 및 행정 연계 계획) 사후 모니터링과 위험도 변화를 플랫폼 이력에 연계",
    }


def _load_font(size: int, bold: bool = False):
    candidates = [
        "/usr/share/fonts/opentype/noto/NotoSansCJK-Bold.ttc" if bold else "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
        "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf" if bold else "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
    ]
    for candidate in candidates:
        if Path(candidate).is_file():
            return ImageFont.truetype(candidate, size=size)
    return ImageFont.load_default()


def _geometry_rings(geometry: dict[str, Any] | None):
    geometry = geometry or {}
    kind = geometry.get("type")
    coords = geometry.get("coordinates") or []
    if kind == "Polygon":
        return coords[:1]
    if kind == "MultiPolygon":
        return [polygon[0] for polygon in coords if polygon]
    return []


def _create_map_image(draft: dict[str, Any], output_path: Path) -> None:
    payload = json.loads(CANDIDATE_GEOJSON.read_text(encoding="utf-8"))
    center_id = str(draft["data_summary"]["center_grid"].get("grid_id", ""))
    neighbor_ids = {str(item.get("grid_id")) for item in draft["data_summary"].get("neighbor_grids", [])}
    features = []
    for feature in payload.get("features", []):
        props = feature.get("properties") or {}
        if _text(_pick(props, "sido")) != draft["sido_name"]:
            continue
        if _text(_pick(props, "sigungu")) != draft["sigungu_name"]:
            continue
        features.append(feature)
    if not features:
        raise ValueError("지도 생성에 사용할 지역 격자가 없습니다.")

    points: list[tuple[float, float]] = []
    for feature in features:
        for ring in _geometry_rings(feature.get("geometry")):
            points.extend((float(x), float(y)) for x, y, *_ in ring)
    min_x = min(x for x, _ in points); max_x = max(x for x, _ in points)
    min_y = min(y for _, y in points); max_y = max(y for _, y in points)
    width, height = 1200, 1240
    margin = 85
    drawable_w = width - margin * 2
    drawable_h = height - 250
    span_x = max(max_x - min_x, 1e-9)
    span_y = max(max_y - min_y, 1e-9)
    scale = min(drawable_w / span_x, drawable_h / span_y)
    offset_x = margin + (drawable_w - span_x * scale) / 2
    offset_y = 140 + (drawable_h - span_y * scale) / 2

    def project(x: float, y: float):
        return (offset_x + (x - min_x) * scale, offset_y + (max_y - y) * scale)

    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    title_font = _load_font(40, True)
    body_font = _load_font(26, False)
    small_font = _load_font(21, False)
    draw.text((60, 40), f"{draft['sido_name']} {draft['sigungu_name']} 신규 확산위험 후보 지도", font=title_font, fill="#0f172a")
    draw.text((60, 94), f"중심 격자 {center_id} · 붉은색 중심 격자 / 주황색 인접 후보격자", font=body_font, fill="#475569")

    for feature in features:
        props = feature.get("properties") or {}
        grid_id = _text(_pick(props, "grid_id"))
        if grid_id == center_id:
            fill, outline, line_width = "#ef4444", "#991b1b", 5
        elif grid_id in neighbor_ids:
            fill, outline, line_width = "#fb923c", "#c2410c", 3
        else:
            fill, outline, line_width = "#dbeafe", "#93c5fd", 1
        for ring in _geometry_rings(feature.get("geometry")):
            polygon = [project(float(x), float(y)) for x, y, *_ in ring]
            if len(polygon) >= 3:
                draw.polygon(polygon, fill=fill, outline=outline, width=line_width)

    legend_y = height - 120
    for x, color, label in [(70, "#ef4444", "중심 격자"), (330, "#fb923c", "인접 후보격자"), (650, "#dbeafe", "지역 내 후보격자")]:
        draw.rectangle((x, legend_y, x + 42, legend_y + 30), fill=color, outline="#64748b")
        draw.text((x + 55, legend_y - 2), label, font=small_font, fill="#334155")
    draw.text((70, height - 66), "※ 감염 확정값이 아닌 신규 발생 후보지역 예측 결과이며 현장 확인이 필요합니다.", font=small_font, fill="#64748b")
    output_path.parent.mkdir(parents=True, exist_ok=True)
    image.save(output_path, format="PNG")


def _replace_docx_media(docx_path: Path, media_name: str, replacement: Path) -> None:
    temp_path = docx_path.with_suffix(".tmp.docx")
    with zipfile.ZipFile(docx_path, "r") as source, zipfile.ZipFile(temp_path, "w", zipfile.ZIP_DEFLATED) as target:
        for item in source.infolist():
            data = replacement.read_bytes() if item.filename == f"word/media/{media_name}" else source.read(item.filename)
            target.writestr(item, data)
    temp_path.replace(docx_path)


def apply_report_template(draft_id: str) -> dict[str, Any]:
    draft = load_draft(draft_id)
    report_type = draft["report_type"]
    template_name = TEMPLATE_FILES.get(report_type)
    if not template_name:
        raise ValueError(f"지원하지 않는 문서 유형입니다: {report_type}")
    template_path = TEMPLATE_ROOT / template_name
    if not template_path.is_file():
        raise FileNotFoundError(f"행정양식 파일이 없습니다: {template_path}")

    output_dir = DATA_ROOT / "generated_drafts" / draft_id / "template"
    output_dir.mkdir(parents=True, exist_ok=True)
    safe_title = re.sub(r'[\\/:*?"<>|]+', "_", draft["title"])
    docx_path = output_dir / f"{safe_title}.docx"
    pdf_path = output_dir / f"{safe_title}.pdf"

    document = Document(template_path)
    replacements = _common_replacements(draft)
    exact = _prediction_exact(draft) if report_type == "prediction" else _field_exact(draft) if report_type == "field_survey" else _control_exact(draft)
    for paragraph in _all_paragraphs(document):
        _replace_paragraph(paragraph, replacements, exact)
    document.save(docx_path)

    if report_type == "prediction":
        map_path = output_dir / "selected_grid_map.png"
        _create_map_image(draft, map_path)
        _replace_docx_media(docx_path, "image1.png", map_path)

    libreoffice = shutil.which("libreoffice") or shutil.which("soffice")
    if not libreoffice:
        raise RuntimeError("LibreOffice가 없어 행정양식 PDF를 생성할 수 없습니다.")
    completed = subprocess.run(
        [libreoffice, "--headless", "--convert-to", "pdf", "--outdir", str(output_dir), str(docx_path)],
        capture_output=True, text=True, timeout=180, check=False,
    )
    if completed.returncode != 0 or not pdf_path.is_file():
        raise RuntimeError(f"행정양식 PDF 변환 실패: {completed.stderr.strip() or completed.stdout.strip()}")

    center = draft["data_summary"].get("center_grid", {})
    template_output = {
        "status": "generated",
        "center_grid_id": center.get("grid_id"),
        "year": draft["year"],
        "sido_name": draft["sido_name"],
        "sigungu_name": draft["sigungu_name"],
        "risk_score": center.get("risk_score"),
        "risk_grade": center.get("risk_grade"),
        "priority_score": center.get("priority_score"),
        "priority_grade": center.get("priority_grade"),
        "docx_path": str(docx_path.resolve()),
        "pdf_path": str(pdf_path.resolve()),
    }
    draft["template_output"] = template_output
    save_draft(draft)
    return template_output


def get_template_file(draft_id: str, file_format: str) -> Path:
    if file_format not in {"docx", "pdf"}:
        raise ValueError("행정양식 파일은 DOCX와 PDF만 지원합니다.")
    draft = load_draft(draft_id)
    output = draft.get("template_output")
    if not isinstance(output, dict):
        raise FileNotFoundError("행정양식이 아직 생성되지 않았습니다.")
    path = Path(str(output.get(f"{file_format}_path", ""))).resolve()
    if not path.is_file():
        raise FileNotFoundError(f"생성 파일을 찾을 수 없습니다: {path}")
    return path


def _next_document_no(csv_path: Path) -> str:
    maximum = 0
    if csv_path.is_file():
        with csv_path.open("r", encoding="utf-8-sig", newline="") as file:
            for row in csv.DictReader(file):
                value = str(row.get("document_no", "")).strip()
                if value.isdigit():
                    maximum = max(maximum, int(value))
    return str(maximum + 1)


def register_report(draft_id: str) -> dict[str, Any]:
    draft = load_draft(draft_id)
    if isinstance(draft.get("registered_report"), dict):
        return draft["registered_report"]

    report_type = draft["report_type"]
    directory = GENERATED_REPORT_ROOT / REPORT_DIRECTORIES[report_type]
    pdf_dir = directory / "pdf"
    docx_dir = directory / "docx"
    pdf_dir.mkdir(parents=True, exist_ok=True)
    docx_dir.mkdir(parents=True, exist_ok=True)
    csv_path = directory / "문서목록.csv"
    document_no = _next_document_no(csv_path)

    source_pdf = get_template_file(draft_id, "pdf")
    source_docx = get_template_file(draft_id, "docx")
    file_stem = f"{document_no}_{draft['year']}_{draft['sigungu_name']}_{REPORT_LABELS[report_type]}"
    pdf_name = f"{file_stem}.pdf"
    docx_name = f"{file_stem}.docx"
    shutil.copy2(source_pdf, pdf_dir / pdf_name)
    shutil.copy2(source_docx, docx_dir / docx_name)

    default_headers = [
        "document_no", "file_name", "year", "center_grid_id", "sido_name", "sigungu_name",
        "risk_score", "risk_grade", "priority_score", "priority_grade", "created_at",
    ]
    existing_headers: list[str] = []
    if csv_path.is_file():
        with csv_path.open("r", encoding="utf-8-sig", newline="") as file:
            existing_headers = next(csv.reader(file), [])
    headers = existing_headers or default_headers
    center = draft["data_summary"].get("center_grid", {})
    row = {
        "document_no": document_no,
        "file_name": pdf_name,
        "year": str(draft["year"]),
        "center_grid_id": str(center.get("grid_id", "")),
        "sido_name": draft["sido_name"],
        "sigungu_name": draft["sigungu_name"],
        "risk_score": _value(center.get("risk_score")),
        "risk_grade": _value(center.get("risk_grade")),
        "priority_score": _value(center.get("priority_score")),
        "priority_grade": _value(center.get("priority_grade")),
        "created_at": datetime.now().astimezone().isoformat(timespec="seconds"),
        "prediction_link_status": "UNLINKED",
        "link_status": "UNLINKED",
        "control_status": "방제 결과 등록",
        "survey_datetime": draft["end_date"],
    }
    write_header = not csv_path.is_file() or csv_path.stat().st_size == 0
    with csv_path.open("a", encoding="utf-8-sig", newline="") as file:
        writer = csv.DictWriter(file, fieldnames=headers, extrasaction="ignore")
        if write_header:
            writer.writeheader()
        writer.writerow({key: row.get(key, "") for key in headers})

    registered = {"report_type": report_type, "document_no": document_no, "file_name": pdf_name}
    draft["status"] = "registered"
    draft["registered_report"] = registered
    save_draft(draft)
    return registered
