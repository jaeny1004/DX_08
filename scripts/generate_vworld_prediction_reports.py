#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
VWorld 배경지도 + 500m 위험격자 + 중심 3x3 예찰권역을 기존 DOCX 양식에 삽입하고
PDF 30건을 생성하는 배치 스크립트.

핵심 원칙
- 기존 발생 예측 보고서 DOCX 양식의 레이아웃/표/별지를 유지한다.
- 문서 본문에 '재구성', '합성', '학습용' 등의 표현을 넣지 않는다.
- 감염 확정 표현 대신 '신규 확산위험 후보', '우선 예찰 검토지역', '현장 확인 필요'를 쓴다.
- 중심 격자 1개를 기준으로 geometry 좌표를 이용해 주변 8개 격자(3x3)를 선택한다.
- VWorld 일반지도 이미지를 받아 위험격자 오버레이를 그린 뒤 템플릿의 지도 이미지를 교체한다.
- 최종 PDF는 LibreOffice headless 변환으로 생성해 원본 양식을 최대한 유지한다.

실행 예시
python scripts/generate_vworld_prediction_reports.py \
  --terrain data/terrain_pine_site_features_south_500m.csv \
  --infection public/data/infection_history_2016_2021.geojson \
  --sigungu data/sigungu_boundary.geojson \
  --template 'rag-backend/data/report_templates/[양식]소나무재선충병 발생 예측 보고서_빈양식.docx' \
  --output rag-backend/data/generated_reports/prediction_30 \
  --count-per-year 5

환경변수(.env)
VWORLD_API_KEY=발급키
VWORLD_API_DOMAIN=101.79.24.212   # VWorld에 등록한 도메인/IP와 동일하게 권장
VWORLD_BASEMAP=GRAPHIC
VWORLD_ZOOM=11
"""

from __future__ import annotations

import argparse
import csv
import io
import json
import math
import os
import re
import shutil
import subprocess
import sys
import time
import zipfile
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Iterable

import geopandas as gpd
import numpy as np
import pandas as pd
import requests
from dotenv import load_dotenv
from PIL import Image, ImageDraw, ImageFont
from pyproj import Transformer
from shapely.geometry import Polygon, box
from docx import Document


# -----------------------------------------------------------------------------
# 기본 설정
# -----------------------------------------------------------------------------

YEAR_RANGE = range(2016, 2022)
GRID_SIZE_M = 500.0
MAP_WIDTH = 1280
MAP_HEIGHT = 880
DEFAULT_ZOOM = 11

RISK_COLORS = {
    "매우 높음": (255, 59, 91, 200),
    "높음": (255, 145, 35, 195),
    "주의": (255, 207, 64, 190),
    "관찰": (57, 197, 105, 180),
    "낮음": (81, 160, 245, 115),
}


@dataclass(frozen=True)
class TerrainCell:
    grid_id: int
    minx: float
    maxy: float
    pine_ratio: float
    elev_mean: float
    slope_mean: float
    site_label: str

    @property
    def geometry(self) -> Polygon:
        return Polygon(
            [
                (self.minx, self.maxy),
                (self.minx + GRID_SIZE_M, self.maxy),
                (self.minx + GRID_SIZE_M, self.maxy - GRID_SIZE_M),
                (self.minx, self.maxy - GRID_SIZE_M),
            ]
        )


@dataclass
class ReportRecord:
    report_no: int
    year: int
    center_grid_id: int
    annual_count: int
    cumulative_count: int
    sido_name: str
    sigungu_name: str


# -----------------------------------------------------------------------------
# 공통 유틸
# -----------------------------------------------------------------------------

def log(message: str) -> None:
    print(message, flush=True)


def sanitize_filename(value: str) -> str:
    return re.sub(r'[\\/:*?"<>|]+', "_", value).strip()


def find_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    candidates = [
        "/usr/share/fonts/truetype/nanum/NanumGothicBold.ttf" if bold else "/usr/share/fonts/truetype/nanum/NanumGothic.ttf",
        "/usr/share/fonts/truetype/nanum/NanumSquareRoundB.ttf" if bold else "/usr/share/fonts/truetype/nanum/NanumSquareRoundR.ttf",
        "/usr/share/fonts/truetype/unfonts-core/UnDotumBold.ttf" if bold else "/usr/share/fonts/truetype/unfonts-core/UnDotum.ttf",
    ]
    for path in candidates:
        if Path(path).exists():
            return ImageFont.truetype(path, size=size)
    return ImageFont.load_default()


def parse_first_wkt_corner(wkt_text: str) -> tuple[float, float]:
    """규칙 격자의 WKT 첫 좌표(minx, maxy)를 빠르게 추출한다."""
    match = re.search(r"POLYGON\s*\(\(\s*([-+0-9.eE]+)\s+([-+0-9.eE]+)", str(wkt_text))
    if not match:
        raise ValueError(f"WKT 시작 좌표를 읽을 수 없습니다: {str(wkt_text)[:100]}")
    return float(match.group(1)), float(match.group(2))


def risk_grade(score: float) -> tuple[int, str]:
    if score >= 85:
        return 1, "매우 높음"
    if score >= 70:
        return 2, "높음"
    if score >= 55:
        return 3, "주의"
    if score >= 40:
        return 4, "관찰"
    return 5, "낮음"


def priority_grade(score: float) -> str:
    if score >= 85:
        return "최우선 예찰"
    if score >= 70:
        return "우선 예찰"
    if score >= 55:
        return "집중 관찰"
    if score >= 40:
        return "정기 관찰"
    return "일반 관리"


# -----------------------------------------------------------------------------
# 데이터 로딩
# -----------------------------------------------------------------------------

def load_terrain_index(csv_path: Path) -> tuple[dict[int, TerrainCell], dict[tuple[float, float], int]]:
    required = {
        "id",
        "WKT",
        "pine_ratio",
        "elev_mean",
        "slope_mean",
        "site_label_mode",
    }
    log(f"[1/8] 지형·소나무 격자 로딩: {csv_path}")

    by_id: dict[int, TerrainCell] = {}
    by_corner: dict[tuple[float, float], int] = {}

    for chunk in pd.read_csv(
        csv_path,
        usecols=lambda col: col in required,
        chunksize=50_000,
        low_memory=False,
    ):
        missing = required - set(chunk.columns)
        if missing:
            raise ValueError(f"terrain CSV 필수 컬럼 누락: {sorted(missing)}")

        for row in chunk.itertuples(index=False):
            grid_id = int(row.id)
            minx, maxy = parse_first_wkt_corner(row.WKT)
            cell = TerrainCell(
                grid_id=grid_id,
                minx=minx,
                maxy=maxy,
                pine_ratio=float(row.pine_ratio) if pd.notna(row.pine_ratio) else 0.0,
                elev_mean=float(row.elev_mean) if pd.notna(row.elev_mean) else 0.0,
                slope_mean=float(row.slope_mean) if pd.notna(row.slope_mean) else 0.0,
                site_label=str(row.site_label_mode) if pd.notna(row.site_label_mode) else "UNKNOWN",
            )
            by_id[grid_id] = cell
            by_corner[(round(minx, 3), round(maxy, 3))] = grid_id

    if not by_id:
        raise RuntimeError("terrain CSV에서 격자를 읽지 못했습니다.")
    return by_id, by_corner


def load_infection_history(path: Path) -> gpd.GeoDataFrame:
    log(f"[2/8] 감염 발생 이력 로딩: {path}")
    gdf = gpd.read_file(path)
    if gdf.crs is None:
        gdf = gdf.set_crs("EPSG:4326")
    gdf = gdf.to_crs("EPSG:5186")

    required = {"id", "infection_count_2016_2021"}
    required.update({f"infection_count_{year}" for year in YEAR_RANGE})
    missing = required - set(gdf.columns)
    if missing:
        raise ValueError(f"infection GeoJSON 필수 속성 누락: {sorted(missing)}")

    gdf["id"] = gdf["id"].astype(int)
    return gdf


def attach_admin_names(infection: gpd.GeoDataFrame, sigungu_path: Path) -> gpd.GeoDataFrame:
    log(f"[3/8] 시군구 경계 결합: {sigungu_path}")
    sigungu = gpd.read_file(sigungu_path)
    if sigungu.crs is None:
        sigungu = sigungu.set_crs("EPSG:4326")
    sigungu = sigungu.to_crs("EPSG:5186")

    required = {"sido_name", "sigungu_name", "geometry"}
    missing = required - set(sigungu.columns)
    if missing:
        raise ValueError(f"sigungu GeoJSON 필수 속성 누락: {sorted(missing)}")

    points = infection[["id", "geometry"]].copy()
    points["geometry"] = points.geometry.centroid
    joined = gpd.sjoin(
        points,
        sigungu[["sido_name", "sigungu_name", "geometry"]],
        how="left",
        predicate="within",
    )
    first = joined.groupby(level=0)[["sido_name", "sigungu_name"]].first()
    result = infection.copy()
    result["sido_name"] = first.reindex(result.index)["sido_name"].fillna("미상").values
    result["sigungu_name"] = first.reindex(result.index)["sigungu_name"].fillna("미상").values
    return result


def choose_report_records(infection: gpd.GeoDataFrame, count_per_year: int) -> list[ReportRecord]:
    records: list[ReportRecord] = []
    report_no = 1
    for year in YEAR_RANGE:
        annual_col = f"infection_count_{year}"
        candidates = infection[infection[annual_col] > 0].copy()
        candidates = candidates.sort_values(
            [annual_col, "infection_count_2016_2021", "id"],
            ascending=[False, False, True],
        ).head(count_per_year)

        for _, row in candidates.iterrows():
            records.append(
                ReportRecord(
                    report_no=report_no,
                    year=year,
                    center_grid_id=int(row["id"]),
                    annual_count=int(row[annual_col]),
                    cumulative_count=int(row["infection_count_2016_2021"]),
                    sido_name=str(row["sido_name"]),
                    sigungu_name=str(row["sigungu_name"]),
                )
            )
            report_no += 1

    expected = len(list(YEAR_RANGE)) * count_per_year
    if len(records) != expected:
        raise RuntimeError(f"보고서 대상이 {expected}건이 아니라 {len(records)}건 선택되었습니다.")
    return records


# -----------------------------------------------------------------------------
# 3x3 격자 및 지표 계산
# -----------------------------------------------------------------------------

def get_3x3_cells(
    center_grid_id: int,
    terrain_by_id: dict[int, TerrainCell],
    terrain_by_corner: dict[tuple[float, float], int],
) -> list[TerrainCell]:
    if center_grid_id not in terrain_by_id:
        raise KeyError(f"terrain CSV에 중심 격자 {center_grid_id}가 없습니다.")

    center = terrain_by_id[center_grid_id]
    cells: list[TerrainCell] = []
    # 위->아래, 왼쪽->오른쪽 순서
    for dy in (GRID_SIZE_M, 0.0, -GRID_SIZE_M):
        for dx in (-GRID_SIZE_M, 0.0, GRID_SIZE_M):
            key = (round(center.minx + dx, 3), round(center.maxy + dy, 3))
            neighbor_id = terrain_by_corner.get(key)
            if neighbor_id is not None:
                cells.append(terrain_by_id[neighbor_id])

    if len(cells) < 4:
        raise RuntimeError(f"격자 {center_grid_id}의 주변 격자가 충분하지 않습니다: {len(cells)}개")
    return cells


def calculate_metrics(
    record: ReportRecord,
    cells: list[TerrainCell],
    infection_by_id: pd.DataFrame,
) -> dict[str, Any]:
    annual_col = f"infection_count_{record.year}"
    block_annual = 0
    block_cumulative = 0
    active_count = 0

    for cell in cells:
        if cell.grid_id in infection_by_id.index:
            row = infection_by_id.loc[cell.grid_id]
            if isinstance(row, pd.DataFrame):
                row = row.iloc[0]
            annual = int(row.get(annual_col, 0))
            cumulative = int(row.get("infection_count_2016_2021", 0))
            block_annual += annual
            block_cumulative += cumulative
            if annual > 0:
                active_count += 1

    pine_mean = float(np.mean([cell.pine_ratio for cell in cells]))
    elevation_mean = float(np.mean([cell.elev_mean for cell in cells]))
    slope_mean = float(np.mean([cell.slope_mean for cell in cells]))

    infection_pressure = min(
        98.0,
        25.0
        + 10.5 * math.log1p(max(block_cumulative, 0))
        + 4.5 * math.log1p(max(block_annual, 0)),
    )
    access_score = max(
        20.0,
        min(92.0, 88.0 - slope_mean * 1.15 - elevation_mean * 0.018),
    )
    road_distance = max(80.0, min(1800.0, 90.0 + slope_mean * 24.0 + elevation_mean * 0.08))

    risk_score = min(
        98.5,
        24.0
        + infection_pressure * 0.39
        + pine_mean * 0.28
        + (100.0 - access_score) * 0.10
        + math.log1p(max(record.annual_count, 1)) * 7.5,
    )
    stage, grade = risk_grade(risk_score)
    priority_score = min(98.0, risk_score * 0.72 + infection_pressure * 0.20 + access_score * 0.08)

    if road_distance < 350:
        road_type = "주요 도로 인접 구간"
    elif road_distance < 800:
        road_type = "일반도로 인접 구간"
    elif road_distance < 1400:
        road_type = "산림 접근 구간"
    else:
        road_type = "산림 내부 접근 구간"

    environment_warning = "해당" if slope_mean >= 25 or elevation_mean >= 450 else "비해당"

    return {
        "block_count": len(cells),
        "block_grid_ids": [cell.grid_id for cell in cells],
        "block_annual": block_annual,
        "block_cumulative": block_cumulative,
        "active_count": active_count,
        "pine_mean": pine_mean,
        "elevation_mean": elevation_mean,
        "slope_mean": slope_mean,
        "infection_pressure": infection_pressure,
        "access_score": access_score,
        "road_distance": road_distance,
        "road_type": road_type,
        "environment_warning": environment_warning,
        "risk_score": risk_score,
        "risk_stage": stage,
        "risk_grade": grade,
        "priority_score": priority_score,
        "priority_grade": priority_grade(priority_score),
    }


# -----------------------------------------------------------------------------
# VWorld 지도 호출 + 오버레이
# -----------------------------------------------------------------------------

def request_vworld_map(
    api_key: str,
    domain: str,
    center_lon: float,
    center_lat: float,
    zoom: int,
    basemap: str,
    width: int = MAP_WIDTH,
    height: int = MAP_HEIGHT,
) -> Image.Image:
    """VWorld 정적 지도 API를 호출한다.

    VWorld 계정/API 설정 차이에 대응할 수 있도록 두 가지 요청 조합을 순서대로 시도한다.
    """
    endpoint = "https://api.vworld.kr/req/image"
    common = {
        "service": "image",
        "request": "GetMap",
        "version": "2.0",
        "key": api_key,
        "domain": domain,
        "format": "png",
        "transparent": "false",
        "center": f"{center_lon:.8f},{center_lat:.8f}",
        "crs": "EPSG:4326",
        "zoom": str(zoom),
        "size": f"{width},{height}",
        "basemap": basemap,
    }

    attempts = [
        common,
        {
            **common,
            "service": "image",
            "request": "getmap",
        },
    ]

    last_error: Exception | None = None
    for params in attempts:
        try:
            response = requests.get(endpoint, params=params, timeout=30)
            response.raise_for_status()
            content_type = response.headers.get("content-type", "").lower()
            if "image" not in content_type:
                message = response.text[:500]
                raise RuntimeError(f"VWorld가 이미지 대신 다음 내용을 반환했습니다: {message}")
            image = Image.open(io.BytesIO(response.content)).convert("RGBA")
            if image.size != (width, height):
                image = image.resize((width, height), Image.Resampling.LANCZOS)
            return image
        except Exception as exc:
            last_error = exc
            time.sleep(1)

    raise RuntimeError(
        "VWorld 지도 이미지를 받지 못했습니다. "
        "VWORLD_API_KEY, VWORLD_API_DOMAIN, 서버 외부 인터넷 연결, VWorld API 승인 상태를 확인하세요. "
        f"마지막 오류: {last_error}"
    )


def lonlat_to_world_pixel(lon: float, lat: float, zoom: int) -> tuple[float, float]:
    lat = max(min(lat, 85.05112878), -85.05112878)
    scale = 256.0 * (2**zoom)
    x = (lon + 180.0) / 360.0 * scale
    sin_lat = math.sin(math.radians(lat))
    y = (0.5 - math.log((1 + sin_lat) / (1 - sin_lat)) / (4 * math.pi)) * scale
    return x, y


def geometry_to_image_points(
    geometry_5186: Polygon,
    center_lon: float,
    center_lat: float,
    zoom: int,
    width: int,
    height: int,
    transformer: Transformer,
) -> list[tuple[float, float]]:
    center_x, center_y = lonlat_to_world_pixel(center_lon, center_lat, zoom)
    points: list[tuple[float, float]] = []
    for x_5186, y_5186 in geometry_5186.exterior.coords:
        lon, lat = transformer.transform(x_5186, y_5186)
        wx, wy = lonlat_to_world_pixel(lon, lat, zoom)
        px = width / 2.0 + (wx - center_x)
        py = height / 2.0 + (wy - center_y)
        points.append((px, py))
    return points


def local_risk_grade(annual: int, cumulative: int) -> str:
    score = 30.0 + 12.0 * math.log1p(max(annual, 0)) + 8.5 * math.log1p(max(cumulative, 0))
    return risk_grade(min(score, 98.0))[1]


def build_vworld_overlay_map(
    output_path: Path,
    record: ReportRecord,
    cells: list[TerrainCell],
    infection: gpd.GeoDataFrame,
    api_key: str,
    domain: str,
    zoom: int,
    basemap: str,
) -> None:
    transformer = Transformer.from_crs("EPSG:5186", "EPSG:4326", always_xy=True)
    center_cell = next(cell for cell in cells if cell.grid_id == record.center_grid_id)
    center_point = center_cell.geometry.centroid
    center_lon, center_lat = transformer.transform(center_point.x, center_point.y)

    image = request_vworld_map(
        api_key=api_key,
        domain=domain,
        center_lon=center_lon,
        center_lat=center_lat,
        zoom=zoom,
        basemap=basemap,
    )
    overlay = Image.new("RGBA", image.size, (0, 0, 0, 0))
    draw = ImageDraw.Draw(overlay, "RGBA")

    # 지도 화면 범위에 들어올 만한 격자만 먼저 공간 필터링한다.
    center_geom = center_cell.geometry
    search_area = center_geom.centroid.buffer(14_000)
    local = infection[infection.intersects(search_area)].copy()
    annual_col = f"infection_count_{record.year}"

    for _, row in local.iterrows():
        annual = int(row.get(annual_col, 0))
        cumulative = int(row.get("infection_count_2016_2021", 0))
        if annual <= 0 and cumulative <= 0:
            continue
        grade = local_risk_grade(annual, cumulative)
        points = geometry_to_image_points(
            row.geometry,
            center_lon,
            center_lat,
            zoom,
            MAP_WIDTH,
            MAP_HEIGHT,
            transformer,
        )
        if points and any(-100 <= x <= MAP_WIDTH + 100 and -100 <= y <= MAP_HEIGHT + 100 for x, y in points):
            draw.polygon(points, fill=RISK_COLORS[grade], outline=(255, 255, 255, 120), width=1)

    # 3x3 주변 격자 강조
    for cell in cells:
        points = geometry_to_image_points(
            cell.geometry,
            center_lon,
            center_lat,
            zoom,
            MAP_WIDTH,
            MAP_HEIGHT,
            transformer,
        )
        if cell.grid_id == record.center_grid_id:
            draw.polygon(points, fill=(255, 56, 56, 75), outline=(160, 0, 0, 255), width=5)
        else:
            draw.line(points, fill=(15, 75, 180, 240), width=4, joint="curve")

    # 3x3 외곽선
    block_union = gpd.GeoSeries([cell.geometry for cell in cells], crs="EPSG:5186").union_all()
    if block_union.geom_type == "Polygon":
        outer_points = geometry_to_image_points(
            block_union,
            center_lon,
            center_lat,
            zoom,
            MAP_WIDTH,
            MAP_HEIGHT,
            transformer,
        )
        draw.line(outer_points, fill=(0, 45, 130, 255), width=6, joint="curve")

    # 중심 격자 라벨
    title_font = find_font(32, bold=True)
    label_font = find_font(24, bold=True)
    small_font = find_font(18, bold=False)

    title = f"{record.sido_name} {record.sigungu_name} / {record.year}년 위험격자 분포"
    title_box = draw.textbbox((0, 0), title, font=title_font)
    title_w = title_box[2] - title_box[0]
    draw.rounded_rectangle(
        (MAP_WIDTH / 2 - title_w / 2 - 18, 16, MAP_WIDTH / 2 + title_w / 2 + 18, 64),
        radius=10,
        fill=(255, 255, 255, 225),
        outline=(160, 160, 160, 220),
        width=2,
    )
    draw.text((MAP_WIDTH / 2 - title_w / 2, 23), title, font=title_font, fill=(20, 20, 20, 255))

    center_label = f"중심 격자 {record.center_grid_id}"
    label_box = draw.textbbox((0, 0), center_label, font=label_font)
    label_w = label_box[2] - label_box[0]
    label_x = MAP_WIDTH / 2 + 20
    label_y = MAP_HEIGHT / 2 - 65
    draw.rounded_rectangle(
        (label_x - 8, label_y - 6, label_x + label_w + 12, label_y + 34),
        radius=8,
        fill=(255, 255, 255, 235),
        outline=(170, 170, 170, 255),
        width=2,
    )
    draw.text((label_x, label_y), center_label, font=label_font, fill=(20, 20, 20, 255))

    # 범례
    legend_x = 24
    legend_y = MAP_HEIGHT - 165
    legend_w = 420
    legend_h = 135
    draw.rounded_rectangle(
        (legend_x, legend_y, legend_x + legend_w, legend_y + legend_h),
        radius=12,
        fill=(255, 255, 255, 235),
        outline=(175, 175, 175, 230),
        width=2,
    )
    legend_items = [
        ("매우 높음", RISK_COLORS["매우 높음"]),
        ("높음", RISK_COLORS["높음"]),
        ("주의", RISK_COLORS["주의"]),
        ("관찰", RISK_COLORS["관찰"]),
    ]
    for idx, (name, color) in enumerate(legend_items):
        col = idx % 2
        row_no = idx // 2
        x = legend_x + 20 + col * 195
        y = legend_y + 17 + row_no * 45
        draw.rectangle((x, y, x + 34, y + 25), fill=color, outline=(255, 255, 255, 220), width=1)
        draw.text((x + 45, y - 1), name, font=small_font, fill=(20, 20, 20, 255))
    draw.rectangle((legend_x + 20, legend_y + 99, legend_x + 54, legend_y + 124), fill=(255, 56, 56, 70), outline=(160, 0, 0, 255), width=4)
    draw.text((legend_x + 65, legend_y + 97), "중심 격자", font=small_font, fill=(20, 20, 20, 255))
    draw.rectangle((legend_x + 215, legend_y + 99, legend_x + 249, legend_y + 124), fill=(255, 255, 255, 30), outline=(15, 75, 180, 255), width=4)
    draw.text((legend_x + 260, legend_y + 97), "주변 8개 격자", font=small_font, fill=(20, 20, 20, 255))

    # 출처 표기
    source = "배경지도: VWorld | 격자: 2016~2021 감염 발생 이력"
    source_box = draw.textbbox((0, 0), source, font=small_font)
    source_w = source_box[2] - source_box[0]
    draw.rounded_rectangle(
        (MAP_WIDTH - source_w - 35, MAP_HEIGHT - 38, MAP_WIDTH - 15, MAP_HEIGHT - 8),
        radius=6,
        fill=(255, 255, 255, 220),
    )
    draw.text((MAP_WIDTH - source_w - 25, MAP_HEIGHT - 35), source, font=small_font, fill=(55, 55, 55, 255))

    result = Image.alpha_composite(image, overlay).convert("RGB")
    output_path.parent.mkdir(parents=True, exist_ok=True)
    result.save(output_path, quality=95)


# -----------------------------------------------------------------------------
# DOCX 템플릿 채우기
# -----------------------------------------------------------------------------

def replace_in_paragraph(paragraph: Any, replacements: dict[str, str]) -> None:
    full = "".join(run.text for run in paragraph.runs)
    if not full:
        return
    new = full
    for old, value in replacements.items():
        new = new.replace(old, str(value))
    if new == full:
        return
    if paragraph.runs:
        paragraph.runs[0].text = new
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.text = new


def replace_everywhere(doc: Document, replacements: dict[str, str]) -> None:
    for paragraph in doc.paragraphs:
        replace_in_paragraph(paragraph, replacements)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_in_paragraph(paragraph, replacements)
    for section in doc.sections:
        for paragraph in section.header.paragraphs:
            replace_in_paragraph(paragraph, replacements)
        for paragraph in section.footer.paragraphs:
            replace_in_paragraph(paragraph, replacements)


def set_paragraph(doc: Document, prefix: str, text: str) -> None:
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith(prefix):
            if paragraph.runs:
                paragraph.runs[0].text = text
                for run in paragraph.runs[1:]:
                    run.text = ""
            else:
                paragraph.text = text
            return
    raise RuntimeError(f"템플릿에서 문단을 찾지 못했습니다: {prefix}")


def replace_docx_media_image(docx_path: Path, image_path: Path, media_name: str = "word/media/image1.png") -> None:
    """원본 DOCX에 포함된 지도 placeholder 이미지를 동일 media 경로로 교체한다.

    이 방식은 원본 문서의 이미지 위치·크기·줄바꿈을 보존한다.
    """
    temp_path = docx_path.with_suffix(".tmp.docx")
    replaced = False
    with zipfile.ZipFile(docx_path, "r") as source, zipfile.ZipFile(
        temp_path, "w", zipfile.ZIP_DEFLATED
    ) as target:
        for item in source.infolist():
            if item.filename == media_name:
                target.writestr(item, image_path.read_bytes())
                replaced = True
            else:
                target.writestr(item, source.read(item.filename))
    if not replaced:
        temp_path.unlink(missing_ok=True)
        raise RuntimeError(f"DOCX 안에서 {media_name}을 찾지 못했습니다.")
    temp_path.replace(docx_path)


def create_docx(
    template_path: Path,
    output_path: Path,
    map_path: Path,
    record: ReportRecord,
    cells: list[TerrainCell],
    metrics: dict[str, Any],
) -> None:
    doc = Document(template_path)
    region = f"{record.sido_name} {record.sigungu_name}"
    grid_ids = metrics["block_grid_ids"]
    adjacent_ids = [grid_id for grid_id in grid_ids if grid_id != record.center_grid_id]
    while len(adjacent_ids) < 4:
        adjacent_ids.append(record.center_grid_id)

    replacements = {
        "[작성일]": f"{record.year}. 12. {10 + (record.report_no % 18):02d}.",
        "-지역, 기간-": f"-{region}, {record.year}년-",
        "[지역]": region,
        "[기간]": f"{record.year}년",
        "[격자 ID]": str(record.center_grid_id),
        "[단계 수]": "5",
        "[단계]": str(metrics["risk_stage"]),
        "[등급]": metrics["risk_grade"],
        "[비율]": f'{metrics["pine_mean"]:.1f}',
        "[거리]": f'{metrics["road_distance"]:.0f}',
        "[도로 유형]": metrics["road_type"],
        "[해당 여부]": metrics["environment_warning"],
        "[기간]": f"{record.year}년",
        "[방향·대상 지역]": "중심 격자 주변 3x3 예찰 검토권역",
        "[결과]": "높음" if metrics["risk_score"] >= 70 else "중간",
        "[방향]": "중심 격자 주변",
        "[인접 격자 1]": str(adjacent_ids[0]),
        "[인접 격자 2]": str(adjacent_ids[1]),
        "[인접 격자 3]": str(adjacent_ids[2]),
        "[인접 격자 4]": str(adjacent_ids[3]),
        "[대응 단계명]": "우선 예찰 검토",
        "[일자]": f"{record.year}. 12. {11 + (record.report_no % 18):02d}.",
        "[대상 격자]": ", ".join(map(str, grid_ids)),
    }
    replace_everywhere(doc, replacements)

    set_paragraph(
        doc,
        "❍ (분석 배경)",
        f"❍ (분석 배경) {record.year}년 감염 발생 이력과 500m 격자 기반 산림·지형정보를 종합하여 신규 확산위험 후보를 분석함",
    )
    set_paragraph(
        doc,
        "❍ (분석 목적)",
        f"❍ (분석 목적) {region} 내 중심 격자 {record.center_grid_id}와 주변 격자를 하나의 예찰 검토권역으로 설정하여 현장 확인 필요성과 예찰 우선순위를 판단",
    )
    set_paragraph(
        doc,
        "❍ (활용 목적)",
        "❍ (활용 목적) AI 위험도 분석 결과를 예찰 계획 수립, 현장 확인, 후속 방제 검토 및 행정 보고에 활용",
    )
    set_paragraph(
        doc,
        "❍ (분석 대상)",
        f"❍ (분석 대상) {region} / {record.year}년 / 중심 격자 포함 3x3 권역 {metrics['block_count']}개 격자",
    )
    set_paragraph(
        doc,
        "❍ (위험 점수)",
        f"❍ (위험 점수) 종합 위험도 스코어 {metrics['risk_score']:.1f}점 (전체 5단계 중 {metrics['risk_stage']}단계 ‘{metrics['risk_grade']}’ 수준)",
    )
    set_paragraph(
        doc,
        "― 최근 감염압력",
        f"― 최근 감염압력 {metrics['infection_pressure']:.1f}점(해당 연도 권역 발생 {metrics['block_annual']}건, 2016~2021 누적 {metrics['block_cumulative']}건 반영)",
    )
    set_paragraph(doc, "― 소나무류 비율", f"― 소나무류 비율 {metrics['pine_mean']:.1f}%")
    set_paragraph(
        doc,
        "❍ 예찰 우선순위",
        f"❍ 예찰 우선순위 {metrics['priority_score']:.1f}점({metrics['priority_grade']})",
    )
    set_paragraph(doc, "❍ 접근성", f"❍ 접근성 {metrics['access_score']:.1f}점")
    set_paragraph(
        doc,
        "― 도로까지의 거리",
        f"― 도로까지의 거리 {metrics['road_distance']:.0f}m({metrics['road_type']})",
    )
    set_paragraph(doc, "― 환경주의", f"― 환경주의 {metrics['environment_warning']}")
    set_paragraph(
        doc,
        "❍ 향후",
        "❍ 향후 3개월 내 중심 격자 주변 3x3 권역으로의 신규 확산위험을 지속적으로 관찰할 필요가 있음",
    )
    set_paragraph(
        doc,
        "❍ (시뮬레이션 종합 의견)",
        f"❍ (시뮬레이션 종합 의견) 중심 격자 {record.center_grid_id}를 포함한 3x3 권역에서 위험 신호가 확인되므로 중심 격자와 인접 격자를 묶어 {metrics['priority_grade']} 대상으로 관리하고 현장 확인 결과를 격자 단위로 기록할 필요가 있다.",
    )
    set_paragraph(doc, "❍ 1단계:", "❍ 1단계: 3x3 권역 우선 예찰")
    set_paragraph(
        doc,
        "― (실행 조건)",
        f"― (실행 조건) 종합 위험도 {metrics['risk_grade']} 또는 감염압력 {metrics['infection_pressure']:.1f}점 이상이며 현장 확인이 완료되지 않은 경우",
    )
    set_paragraph(
        doc,
        "― (실행 계획)",
        f"― (실행 계획) 중심 격자 {record.center_grid_id} 및 주변 {metrics['block_count'] - 1}개 격자를 대상으로 현장 예찰 또는 드론 예찰 시행",
    )
    set_paragraph(doc, "❍ 2단계:", "❍ 2단계: 현장 확인 및 검경 연계")
    set_paragraph(
        doc,
        "― (대상 조직·알림 내용)",
        f"― (대상 조직·알림 내용) 관할 산림부서에 {region} 3x3 예찰 검토권역의 위험도, 감염 발생 이력 및 예찰 우선순위를 공유",
    )
    set_paragraph(
        doc,
        "― (현장 확인 계획)",
        "― (현장 확인 계획) 변색목·고사목 유무를 확인하고 의심목 발견 시 시료 채취 및 전문기관 검경 의뢰",
    )
    set_paragraph(doc, "❍ 3단계:", "❍ 3단계: 방제 검토 및 사후 모니터링")
    set_paragraph(
        doc,
        "― (방제·수종전환 검토 사항)",
        "― (방제·수종전환 검토 사항) 현장 확인 결과와 주변 소나무류 분포를 확인한 뒤 방제 대상 지정 및 임분 관리 필요성을 검토",
    )
    set_paragraph(
        doc,
        "― (행정 연계 및 후속 계획)",
        "― (행정 연계 및 후속 계획) 예찰 결과를 격자 단위 이력으로 등록하고 인접 권역 모니터링 주기를 조정",
    )

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    replace_docx_media_image(output_path, map_path)


# -----------------------------------------------------------------------------
# PDF 변환 및 검증
# -----------------------------------------------------------------------------

def find_libreoffice() -> str:
    for name in ("libreoffice", "soffice"):
        path = shutil.which(name)
        if path:
            return path
    raise RuntimeError("LibreOffice가 설치되어 있지 않습니다. sudo apt install libreoffice 로 설치하세요.")


def convert_docx_to_pdf(docx_path: Path, pdf_dir: Path) -> Path:
    libreoffice = find_libreoffice()
    pdf_dir.mkdir(parents=True, exist_ok=True)
    profile_dir = pdf_dir / ".lo_profile"
    profile_dir.mkdir(parents=True, exist_ok=True)

    command = [
        libreoffice,
        "--headless",
        f"-env:UserInstallation=file://{profile_dir.resolve()}",
        "--convert-to",
        "pdf:writer_pdf_Export",
        "--outdir",
        str(pdf_dir),
        str(docx_path),
    ]
    result = subprocess.run(command, text=True, capture_output=True, timeout=180)
    if result.returncode != 0:
        raise RuntimeError(
            f"PDF 변환 실패: {docx_path.name}\nSTDOUT: {result.stdout}\nSTDERR: {result.stderr}"
        )

    pdf_path = pdf_dir / f"{docx_path.stem}.pdf"
    if not pdf_path.exists() or pdf_path.stat().st_size == 0:
        raise RuntimeError(f"PDF 파일이 생성되지 않았습니다: {pdf_path}")
    return pdf_path


def verify_pdf_count(pdf_dir: Path, expected: int) -> None:
    pdfs = list(pdf_dir.glob("*.pdf"))
    if len(pdfs) != expected:
        raise RuntimeError(f"PDF가 {expected}개가 아니라 {len(pdfs)}개 생성되었습니다.")
    for pdf in pdfs:
        if pdf.stat().st_size < 10_000:
            raise RuntimeError(f"PDF 크기가 비정상적으로 작습니다: {pdf.name}")


# -----------------------------------------------------------------------------
# 메인
# -----------------------------------------------------------------------------

def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(description="기존 양식 기반 VWorld 발생 예측 PDF 30건 생성")
    parser.add_argument("--terrain", type=Path, required=True)
    parser.add_argument("--infection", type=Path, required=True)
    parser.add_argument("--sigungu", type=Path, required=True)
    parser.add_argument("--template", type=Path, required=True)
    parser.add_argument("--output", type=Path, required=True)
    parser.add_argument("--count-per-year", type=int, default=5)
    parser.add_argument("--zoom", type=int, default=None)
    parser.add_argument("--keep-docx", action="store_true", help="중간 DOCX도 보관")
    return parser


def main() -> int:
    args = build_parser().parse_args()
    load_dotenv(Path(__file__).resolve().parents[1] / "rag-backend" / ".env")
    load_dotenv()  # 현재 작업 폴더의 .env도 허용

    api_key = os.getenv("VWORLD_API_KEY", "").strip()
    domain = os.getenv("VWORLD_API_DOMAIN", "").strip()
    basemap = os.getenv("VWORLD_BASEMAP", "GRAPHIC").strip() or "GRAPHIC"
    zoom = args.zoom or int(os.getenv("VWORLD_ZOOM", str(DEFAULT_ZOOM)))

    if not api_key:
        raise RuntimeError(".env에 VWORLD_API_KEY가 없습니다.")
    if not domain:
        raise RuntimeError(".env에 VWORLD_API_DOMAIN이 없습니다. VWorld에 등록한 서버 도메인 또는 IP를 입력하세요.")

    for path in (args.terrain, args.infection, args.sigungu, args.template):
        if not path.exists():
            raise FileNotFoundError(path)

    output_root = args.output.resolve()
    docx_dir = output_root / "docx"
    pdf_dir = output_root / "pdf"
    map_dir = output_root / "maps"
    output_root.mkdir(parents=True, exist_ok=True)
    docx_dir.mkdir(exist_ok=True)
    pdf_dir.mkdir(exist_ok=True)
    map_dir.mkdir(exist_ok=True)

    terrain_by_id, terrain_by_corner = load_terrain_index(args.terrain)
    infection = load_infection_history(args.infection)
    infection = attach_admin_names(infection, args.sigungu)
    records = choose_report_records(infection, args.count_per_year)
    infection_by_id = infection.set_index("id", drop=False)

    manifest: list[dict[str, Any]] = []
    total = len(records)

    log(f"[4/8] 총 {total}건 지도·문서 생성 시작")
    for record in records:
        log(f"  - ({record.report_no:02d}/{total}) {record.year}년 격자 {record.center_grid_id}")
        cells = get_3x3_cells(record.center_grid_id, terrain_by_id, terrain_by_corner)
        metrics = calculate_metrics(record, cells, infection_by_id)

        base_name = sanitize_filename(
            f"{record.report_no:02d}_{record.year}_소나무재선충병_발생예측보고서_"
            f"{record.sido_name}_{record.sigungu_name}_격자{record.center_grid_id}"
        )
        map_path = map_dir / f"{base_name}.png"
        docx_path = docx_dir / f"{base_name}.docx"

        build_vworld_overlay_map(
            output_path=map_path,
            record=record,
            cells=cells,
            infection=infection,
            api_key=api_key,
            domain=domain,
            zoom=zoom,
            basemap=basemap,
        )
        create_docx(
            template_path=args.template,
            output_path=docx_path,
            map_path=map_path,
            record=record,
            cells=cells,
            metrics=metrics,
        )
        pdf_path = convert_docx_to_pdf(docx_path, pdf_dir)

        manifest.append(
            {
                "document_no": record.report_no,
                "file_name": pdf_path.name,
                "year": record.year,
                "center_grid_id": record.center_grid_id,
                "sido_name": record.sido_name,
                "sigungu_name": record.sigungu_name,
                "block_grid_ids": "|".join(map(str, metrics["block_grid_ids"])),
                "center_annual_count": record.annual_count,
                "center_cumulative_count": record.cumulative_count,
                "block_annual_count": metrics["block_annual"],
                "block_cumulative_count": metrics["block_cumulative"],
                "risk_score": round(metrics["risk_score"], 1),
                "risk_grade": metrics["risk_grade"],
                "priority_score": round(metrics["priority_score"], 1),
                "priority_grade": metrics["priority_grade"],
            }
        )

    log("[5/8] PDF 개수 및 파일 크기 검증")
    verify_pdf_count(pdf_dir, total)

    manifest_path = output_root / "문서목록.csv"
    pd.DataFrame(manifest).to_csv(manifest_path, index=False, encoding="utf-8-sig")

    log("[6/8] PDF ZIP 생성")
    zip_path = output_root / "발생예측보고서_30건_VWorld_기존양식_PDF.zip"
    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as archive:
        for pdf in sorted(pdf_dir.glob("*.pdf")):
            archive.write(pdf, arcname=pdf.name)
        archive.write(manifest_path, arcname=manifest_path.name)

    if not args.keep_docx:
        shutil.rmtree(docx_dir, ignore_errors=True)

    log("[7/8] 완료")
    log(f"PDF 폴더: {pdf_dir}")
    log(f"ZIP 파일: {zip_path}")
    log("[8/8] 처음에는 PDF 1~2개를 열어 지도 위치와 양식 줄바꿈을 확인하세요.")
    return 0


if __name__ == "__main__":
    try:
        raise SystemExit(main())
    except KeyboardInterrupt:
        print("사용자에 의해 중단되었습니다.", file=sys.stderr)
        raise SystemExit(130)
    except Exception as exc:
        print(f"오류: {exc}", file=sys.stderr)
        raise SystemExit(1)
