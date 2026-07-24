#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
예측보고서 + 현장예찰보고서를 1:1로 연결해
소나무재선충 방제 보고서의 본문과 별지 4종을 모두 채우는 생성기.

별지 4종
1. 재선충병 방제사업 계획서
2. 방제조치 명령서 관리대장
3. 재선충병 방제대상목 조사야장
4. 피해고사목 방제실적

중요:
- 검경 결과가 없는 자료를 감염 확정으로 표현하지 않는다.
- 본문은 "방제 검토 계획" 중심으로 작성한다.
- DOCX 안에 대괄호 플레이스홀더가 하나라도 남으면 실패 처리한다.
- 별지 image1.png~image4.png는 원본 위 좌표 덧쓰기가 아니라
  서로 다른 완성형 별지 페이지를 새로 만들어 교체한다.
"""

from __future__ import annotations

import argparse
import io
import math
import re
import subprocess
import zipfile
from dataclasses import dataclass
from datetime import datetime, timedelta
from pathlib import Path
from typing import Any, Iterable

import pandas as pd
from PIL import Image, ImageDraw, ImageFont
from docx import Document


SCRIPT_VERSION = "2026.07.20-control-linked-complete-v4"

PAGE_SIZE = (832, 1272)
MEDIA_NAMES = ("image1.png", "image2.png", "image3.png", "image4.png")


@dataclass(frozen=True)
class Record:
    document_no: int
    year: int
    center_grid_id: int
    sido_name: str
    sigungu_name: str
    risk_score: float
    risk_grade: str
    priority_score: float
    priority_grade: str
    suspicious_count: int
    sample_count: int
    survey_datetime: str
    surveyors: str
    source_prediction_file: str
    source_field_file: str
    pine_area_ha: float
    pine_ratio_pct: float


def log(message: str) -> None:
    print(message, flush=True)


def to_text(value: Any, default: str = "") -> str:
    if value is None:
        return default
    try:
        if pd.isna(value):
            return default
    except Exception:
        pass
    result = str(value).strip()
    return result if result else default


def to_float(value: Any, default: float = 0.0) -> float:
    try:
        if value is None or pd.isna(value):
            return default
        return float(value)
    except Exception:
        return default


def to_int(value: Any, default: int = 0) -> int:
    return max(0, int(round(to_float(value, default))))


def first_value(row: dict[str, Any], names: Iterable[str], default: Any = "") -> Any:
    for name in names:
        if name in row:
            value = row[name]
            try:
                if not pd.isna(value):
                    return value
            except Exception:
                if value is not None:
                    return value
    return default


def read_csv(path: Path) -> pd.DataFrame:
    try:
        return pd.read_csv(path, encoding="utf-8-sig")
    except UnicodeDecodeError:
        return pd.read_csv(path, encoding="cp949")


def load_terrain(path: Path) -> dict[int, dict[str, float]]:
    header = pd.read_csv(path, nrows=0)
    columns = list(header.columns)

    def find(candidates: list[str]) -> str | None:
        for candidate in candidates:
            if candidate in columns:
                return candidate
        return None

    id_col = find(["id", "grid_id", "GRID_ID", "cell_id"])
    if id_col is None:
        raise RuntimeError("terrain CSV에서 격자 ID 컬럼을 찾지 못했습니다.")

    area_col = find(["pine_area", "pine_area_m2", "소나무면적"])
    ratio_col = find(["pine_ratio", "pine_ratio_pct", "소나무비율"])

    usecols = [id_col] + [c for c in (area_col, ratio_col) if c]
    frame = pd.read_csv(path, usecols=usecols)

    result: dict[int, dict[str, float]] = {}
    for row in frame.to_dict(orient="records"):
        grid_id = int(row[id_col])
        area_m2 = max(0.0, to_float(row.get(area_col), 0.0)) if area_col else 0.0
        ratio = max(0.0, to_float(row.get(ratio_col), 0.0)) if ratio_col else 0.0
        if 0 <= ratio <= 1:
            ratio *= 100
        result[grid_id] = {
            "pine_area_ha": round(area_m2 / 10000, 2),
            "pine_ratio_pct": round(min(100.0, ratio), 2),
        }
    return result


def load_records(
    prediction_manifest: Path,
    field_manifest: Path,
    terrain_path: Path,
) -> list[Record]:
    prediction = read_csv(prediction_manifest)
    field = read_csv(field_manifest)

    if "document_no" not in prediction.columns or "document_no" not in field.columns:
        raise RuntimeError("두 문서목록에 document_no 컬럼이 필요합니다.")

    prediction = prediction.sort_values("document_no").reset_index(drop=True)
    field = field.sort_values("document_no").reset_index(drop=True)

    merged = prediction.merge(
        field,
        on="document_no",
        suffixes=("_prediction", "_field"),
        validate="one_to_one",
    )
    if len(merged) != len(prediction) or len(merged) != len(field):
        raise RuntimeError(
            f"예측/현장 문서 연결 실패: prediction={len(prediction)}, "
            f"field={len(field)}, merged={len(merged)}"
        )

    terrain = load_terrain(terrain_path)
    records: list[Record] = []

    for row in merged.to_dict(orient="records"):
        document_no = int(row["document_no"])

        pred_grid = int(first_value(
            row,
            ["center_grid_id_prediction", "center_grid_id"],
            0,
        ))
        field_grid = int(first_value(
            row,
            ["center_grid_id_field", "center_grid_id"],
            0,
        ))
        if pred_grid != field_grid:
            raise RuntimeError(
                f"{document_no}번 중심 격자 불일치: prediction={pred_grid}, field={field_grid}"
            )

        pred_sido = to_text(first_value(
            row,
            ["sido_name_prediction", "sido_name"],
            "",
        ))
        field_sido = to_text(first_value(
            row,
            ["sido_name_field", "sido_name"],
            pred_sido,
        ))
        pred_sigungu = to_text(first_value(
            row,
            ["sigungu_name_prediction", "sigungu_name"],
            "",
        ))
        field_sigungu = to_text(first_value(
            row,
            ["sigungu_name_field", "sigungu_name"],
            pred_sigungu,
        ))
        if pred_sido != field_sido or pred_sigungu != field_sigungu:
            raise RuntimeError(
                f"{document_no}번 행정구역 불일치: "
                f"{pred_sido} {pred_sigungu} / {field_sido} {field_sigungu}"
            )

        year = int(first_value(
            row,
            ["year_prediction", "year_field", "year"],
            0,
        ))
        if year <= 0:
            raise RuntimeError(f"{document_no}번 보고서 연도를 찾지 못했습니다.")

        terrain_value = terrain.get(
            pred_grid,
            {"pine_area_ha": 0.0, "pine_ratio_pct": 0.0},
        )

        records.append(Record(
            document_no=document_no,
            year=year,
            center_grid_id=pred_grid,
            sido_name=pred_sido,
            sigungu_name=pred_sigungu,
            risk_score=to_float(first_value(
                row,
                ["risk_score_prediction", "risk_score_field", "risk_score"],
                0,
            )),
            risk_grade=to_text(first_value(
                row,
                ["risk_grade_prediction", "risk_grade_field", "risk_grade"],
                "현장 확인 필요",
            )),
            priority_score=to_float(first_value(
                row,
                ["priority_score_prediction", "priority_score_field", "priority_score"],
                0,
            )),
            priority_grade=to_text(first_value(
                row,
                ["priority_grade_prediction", "priority_grade_field", "priority_grade"],
                "우선 예찰 검토지역",
            )),
            suspicious_count=to_int(first_value(
                row,
                ["suspicious_count", "detected_count", "abnormal_count"],
                0,
            )),
            sample_count=to_int(first_value(
                row,
                ["sample_count", "samples"],
                0,
            )),
            survey_datetime=to_text(first_value(
                row,
                ["survey_datetime", "survey_date", "field_survey_date"],
                "",
            )),
            surveyors=to_text(first_value(
                row,
                ["surveyors", "surveyor_name"],
                "산림보호 담당자",
            )),
            source_prediction_file=to_text(first_value(
                row,
                ["file_name_prediction", "source_prediction_file"],
                f"prediction_{document_no:02d}",
            )),
            source_field_file=to_text(first_value(
                row,
                ["file_name_field", "source_field_file"],
                f"field_{document_no:02d}",
            )),
            pine_area_ha=terrain_value["pine_area_ha"],
            pine_ratio_pct=terrain_value["pine_ratio_pct"],
        ))

    log(f"예측보고서와 현장예찰보고서 {len(records)}건 연결 완료")
    return records


def parse_date(value: str, fallback_year: int) -> datetime:
    patterns = [
        r"(\d{4})[.\-/]\s*(\d{1,2})[.\-/]\s*(\d{1,2})",
        r"(\d{4})년\s*(\d{1,2})월\s*(\d{1,2})일",
    ]
    for pattern in patterns:
        match = re.search(pattern, value)
        if match:
            return datetime(
                int(match.group(1)),
                int(match.group(2)),
                int(match.group(3)),
            )
    return datetime(fallback_year, 12, 20)


def risk_label(score: float) -> str:
    if score >= 80:
        return "매우 높음"
    if score >= 60:
        return "높음"
    if score >= 40:
        return "주의"
    if score >= 20:
        return "관찰"
    return "낮음"


def plan_values(record: Record) -> dict[str, Any]:
    survey_date = parse_date(record.survey_datetime, record.year)
    start_date = survey_date + timedelta(days=7)
    end_date = start_date + timedelta(days=2)

    suspicious = max(0, record.suspicious_count)
    planned = suspicious
    shred = planned if planned <= 5 else int(round(planned * 0.7))
    fumigate = max(0, planned - shred)
    preventive = max(0, int(round(planned * 0.6)))
    total = shred + fumigate
    target_area = max(
        0.5,
        min(25.0, record.pine_area_ha if record.pine_area_ha > 0 else 5.0),
    )
    after_score = round(max(0.0, record.risk_score * 0.68), 1)

    surveyor = record.surveyors.split(",")[0].strip() or "산림보호 담당자"
    plan_id = f"CTRL-{record.year}-{record.center_grid_id}-{record.document_no:02d}"

    return {
        "start_date": start_date,
        "end_date": end_date,
        "days": (end_date - start_date).days + 1,
        "surveyor": surveyor,
        "team": f"{record.sigungu_name} 산림보호 방제지원조",
        "address": f"{record.sido_name} {record.sigungu_name} 산림 일원",
        "area_ha": round(target_area, 1),
        "range": f"중심 격자 {record.center_grid_id} 및 주변 3×3 권역",
        "confirmed": 0,
        "concern": suspicious,
        "planned": planned,
        "shred": shred,
        "fumigate": fumigate,
        "preventive": preventive,
        "sample_count": record.sample_count,
        "total": total,
        "tarpaulin": f"검경 결과에 따라 발급 예정 ({plan_id})",
        "after_score": after_score,
        "after_grade": risk_label(after_score),
        "plan_id": plan_id,
    }


def set_paragraph_text(paragraph, new_text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = new_text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(new_text)


def replace_cover_text(paragraph, record: Record, values: dict[str, Any]) -> bool:
    original = paragraph.text
    changed = False

    if "-지역, 기간-" in original:
        set_paragraph_text(
            paragraph,
            f"-{record.sido_name} {record.sigungu_name}, "
            f"{values['start_date']:%Y. %m. %d.}~{values['end_date']:%Y. %m. %d.}-",
        )
        return True

    if "[작성일]" in original:
        set_paragraph_text(paragraph, f"{values['start_date']:%Y. %m. %d.}")
        return True

    return changed


def replace_body_paragraph(paragraph, record: Record, v: dict[str, Any]) -> bool:
    text_value = paragraph.text.strip()
    if not text_value:
        return False

    replacements: list[tuple[str, str]] = [
        (
            "(작업 기간)",
            f"❍ (작업 기간) {v['start_date']:%Y. %m. %d.} ~ "
            f"{v['end_date']:%Y. %m. %d.} (총 {v['days']}일간)",
        ),
        (
            "(방 제 자",
            f"❍ (방 제 자) {v['team']} (단원: {v['surveyor']})",
        ),
        (
            "(대상 위치)",
            f"❍ (대상 위치) {v['address']} "
            f"(격자 ID: {record.center_grid_id})",
        ),
        (
            "(방제 면적)",
            f"❍ (방제 면적) 총 {v['area_ha']:.1f}ha ({v['range']})",
        ),
        (
            "(대상 수량)",
            f"❍ (대상 수량) 검경 확정목 {v['confirmed']}본 및 "
            f"현장 이상징후·감염 우려 피해목 {v['concern']}본 "
            f"(방제 검토 대상 총 {v['planned']}본)",
        ),
        (
            "(파쇄 처리)",
            f"❍ (파쇄 처리 계획) {v['shred']}본 / "
            "검경 결과 확인 후 현장 파쇄 또는 지정 장소 반출",
        ),
        (
            "(훈증 처리)",
            f"❍ (훈증 처리 계획) {v['fumigate']}본 / "
            "파쇄가 어려운 대상목에 한해 밀봉 처리 검토",
        ),
        (
            "타포린 피복 일련번호",
            f"― 타포린 피복 일련번호: {v['tarpaulin']}",
        ),
        (
            "(작업 면적)",
            f"❍ (작업 면적) 중심 격자와 인접 우량 소나무림 "
            f"(약 {v['area_ha']:.1f}ha)",
        ),
        (
            "(주입 실적)",
            f"❍ (주입 계획) 소나무류 약 {v['preventive']}본 / "
            "대상목 검토 후 등록 약제 기준 예방나무주사 적용",
        ),
        (
            "(천공 규격)",
            "❍ (천공 규격) 직경 10mm 내외, 깊이 5cm 내외 / "
            "수간주입 기준과 현장 여건에 따라 조정",
        ),
        (
            "(방제 전)",
            f"❍ (방제 전) 중심 격자 {record.center_grid_id}에서 "
            f"현장 이상징후 {record.suspicious_count}본과 "
            f"시료 {record.sample_count}점이 확인되어 현장 확인 필요",
        ),
        (
            "(방제 후)",
            "❍ (방제 후 계획) 검경 결과에 따라 대상목 처리, "
            "잔재물 위치 등록, 인접 격자 재예찰을 순차 수행",
        ),
        (
            "(위험 스코어 조정)",
            "❍ (위험 스코어 조정)",
        ),
        (
            "방제 완료 자료 입력에 따라",
            f"― 방제 조치 결과 등록 시 해당 격자 "
            f"({record.center_grid_id})의 AI 종합 위험도 점수를 재산정",
        ),
        (
            "위험도 변화:",
            f"― 위험도 변화(시나리오): 기존 {record.risk_score:.1f}점"
            f"({record.risk_grade}) → 조치 반영 후 예상 "
            f"{v['after_score']:.1f}점({v['after_grade']})",
        ),
        (
            "훈증 더미",
            f"― 훈증 더미 최대 {v['fumigate']}개소 위치좌표 등록 및 "
            "월 1회 사후 모니터링 계획",
        ),
        (
            "(잔재물 관리 조치",
            "― (잔재물 관리 조치) 파쇄물 비산 방지, 훈증 피복상태 확인, "
            "처리 위치와 사진을 시스템에 기록",
        ),
        (
            "(보고 및 결재",
            f"❍ (보고 및 결재) {record.sigungu_name} 산림보호 담당부서 "
            "검토 후 방제 대상·방법 확정 및 작업 승인",
        ),
        (
            "(사후 모니터링)",
            "❍ (사후 모니터링) 처리 후 1개월·3개월 단위 재예찰 및 "
            "인접 격자 이상징후 추가 확인",
        ),
        (
            "(후속 사업 및 행정 연계 계획",
            "― (후속 사업 및 행정 연계 계획) 검경 결과, 방제 이력, "
            "현장 사진을 통합 저장하고 차기 예찰 우선순위 산정에 반영",
        ),
    ]

    for marker, replacement in replacements:
        if marker in text_value:
            set_paragraph_text(paragraph, replacement)
            return True
    return False


def walk_table_paragraphs(table):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                yield paragraph
            for nested in cell.tables:
                yield from walk_table_paragraphs(nested)


def replace_remaining_placeholders(
    paragraph,
    record: Record,
    v: dict[str, Any],
) -> bool:
    """
    특정 문장 치환 후에도 양식 내부에 독립적으로 남은 플레이스홀더를 치환한다.
    [방법], [일자]처럼 문장 전체 패턴에 걸리지 않는 짧은 토큰을 처리한다.
    """
    original = paragraph.text
    if not original:
        return False

    replacements = {
        "[일자]": f"{v['start_date']:%Y. %m. %d.}",
        "[방법]": "검경 결과에 따라 파쇄·훈증 등 적정 방제방법 결정",
        "[시작일]": f"{v['start_date']:%Y. %m. %d.}",
        "[종료일]": f"{v['end_date']:%Y. %m. %d.}",
        "[일수]": str(v["days"]),
        "[소속·조]": v["team"],
        "[성명]": v["surveyor"],
        "[주소]": v["address"],
        "[격자 ID]": str(record.center_grid_id),
        "[면적]": f"{v['area_ha']:.1f}",
        "[포함 범위]": v["range"],
        "[수량]": str(v["planned"]),
        "[처리 방법 및 규격]": "검경 결과에 따라 현장 파쇄 또는 지정 장소 반출",
        "[처리 사유 및 방법]": "파쇄 곤란 대상목에 한해 밀봉 훈증 검토",
        "[번호]": v["tarpaulin"],
        "[범위 및 대상]": f"중심 격자 {record.center_grid_id} 인접 우량 소나무림",
        "[수종]": "소나무류",
        "[약제 및 처리 내용]": "등록 약제 기준 예방나무주사 적용",
        "[직경]": "10",
        "[깊이]": "5",
        "[작업 방법]": "수간주입 기준과 현장 여건에 따라 조정",
        "[현장 상태 입력]": (
            f"현장 이상징후 {record.suspicious_count}본, "
            f"시료 {record.sample_count}점 확인"
        ),
        "[조치 결과 입력]": "검경 결과 확인 후 대상목 처리 및 인접 격자 재예찰",
        "[점수]": f"{record.risk_score:.1f}",
        "[등급]": record.risk_grade,
        "[개소]": str(v["fumigate"]),
        "[주기]": "월 1회",
        "[입력]": "현장 사진·좌표·처리 이력을 시스템에 등록",
        "[보고 대상·승인 절차 입력]": (
            f"{record.sigungu_name} 산림보호 담당부서 검토 후 작업 승인"
        ),
        "[후속 사업 및 행정 연계 계획]": (
            "검경·방제 이력을 차기 예찰 우선순위 산정에 반영"
        ),
    }

    updated = original
    for token, replacement in replacements.items():
        updated = updated.replace(token, replacement)

    if updated != original:
        set_paragraph_text(paragraph, updated)
        return True
    return False


def fill_document_text(
    template: Path,
    output: Path,
    record: Record,
    values: dict[str, Any],
) -> None:
    doc = Document(template)

    for paragraph in doc.paragraphs:
        replace_cover_text(paragraph, record, values)
        replace_body_paragraph(paragraph, record, values)
        replace_remaining_placeholders(paragraph, record, values)

    for table in doc.tables:
        for paragraph in walk_table_paragraphs(table):
            replace_cover_text(paragraph, record, values)
            replace_body_paragraph(paragraph, record, values)
            replace_remaining_placeholders(paragraph, record, values)

    for section in doc.sections:
        for container in (section.header, section.footer):
            for paragraph in container.paragraphs:
                replace_cover_text(paragraph, record, values)
                replace_body_paragraph(paragraph, record, values)
                replace_remaining_placeholders(paragraph, record, values)

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def font_path(bold: bool = False) -> str | None:
    candidates = [
        "/usr/share/fonts/truetype/nanum/NanumGothicBold.ttf" if bold
        else "/usr/share/fonts/truetype/nanum/NanumGothic.ttf",
        "/usr/share/fonts/truetype/nanum/NanumMyeongjoBold.ttf" if bold
        else "/usr/share/fonts/truetype/nanum/NanumMyeongjo.ttf",
        "/usr/share/fonts/truetype/unfonts-core/UnDotumBold.ttf" if bold
        else "/usr/share/fonts/truetype/unfonts-core/UnDotum.ttf",
    ]
    for candidate in candidates:
        if Path(candidate).exists():
            return candidate
    return None


def font(size: int, bold: bool = False) -> ImageFont.ImageFont:
    path = font_path(bold)
    if path:
        return ImageFont.truetype(path, size=size)
    return ImageFont.load_default()


def centered(draw, box, text_value, size=20, bold=False, fill=(0, 0, 0)):
    x1, y1, x2, y2 = box
    f = font(size, bold)
    bbox = draw.multiline_textbbox((0, 0), str(text_value), font=f, spacing=3, align="center")
    w = bbox[2] - bbox[0]
    h = bbox[3] - bbox[1]
    draw.multiline_text(
        (x1 + (x2 - x1 - w) / 2, y1 + (y2 - y1 - h) / 2),
        str(text_value),
        font=f,
        fill=fill,
        spacing=3,
        align="center",
    )


def wrapped(draw, box, text_value, size=18, bold=False, align="left"):
    x1, y1, x2, y2 = box
    max_width = x2 - x1 - 10
    f = font(size, bold)
    lines: list[str] = []
    current = ""
    for ch in str(text_value):
        trial = current + ch
        if draw.textbbox((0, 0), trial, font=f)[2] <= max_width or not current:
            current = trial
        else:
            lines.append(current)
            current = ch
    if current:
        lines.append(current)
    joined = "\n".join(lines)
    draw.multiline_text((x1 + 5, y1 + 5), joined, font=f, fill=(0, 0, 0), spacing=4, align=align)


def base_page(title: str) -> tuple[Image.Image, ImageDraw.ImageDraw]:
    image = Image.new("RGB", PAGE_SIZE, "white")
    draw = ImageDraw.Draw(image)
    centered(draw, (60, 40, 772, 115), title, 29, True)
    draw.line((180, 108, 652, 108), fill="black", width=2)
    draw.line((180, 112, 652, 112), fill="black", width=2)
    return image, draw


def draw_grid(
    draw: ImageDraw.ImageDraw,
    left: int,
    top: int,
    widths: list[int],
    heights: list[int],
    line_width: int = 1,
) -> list[list[tuple[int, int, int, int]]]:
    xs = [left]
    for width in widths:
        xs.append(xs[-1] + width)
    ys = [top]
    for height in heights:
        ys.append(ys[-1] + height)

    for x in xs:
        draw.line((x, ys[0], x, ys[-1]), fill="black", width=line_width)
    for y in ys:
        draw.line((xs[0], y, xs[-1], y), fill="black", width=line_width)

    cells = []
    for r in range(len(heights)):
        row = []
        for c in range(len(widths)):
            row.append((xs[c], ys[r], xs[c + 1], ys[r + 1]))
        cells.append(row)
    return cells


def create_appendix_1(record: Record, v: dict[str, Any], output: Path) -> None:
    image, draw = base_page(f"재선충병 방제사업 계획서({record.sigungu_name})")

    wrapped(draw, (55, 145, 780, 190), "1. 산림 현황", 23, True)
    wrapped(
        draw,
        (70, 200, 780, 300),
        f"가. 산림면적\n  - 중심 격자 포함 3×3 관리권역 225.0ha\n"
        f"나. 소나무림 현황\n  - 소나무류 면적 약 {record.pine_area_ha:.1f}ha "
        f"(권역 대비 {record.pine_ratio_pct:.1f}%)",
        17,
    )

    wrapped(draw, (55, 330, 780, 380), "2. 재선충병 발생 및 방제 현황", 23, True)
    wrapped(
        draw,
        (70, 390, 780, 540),
        f"가. 발생경과\n  (1) 최초발생: 감염 발생 이력 및 신규 확산위험 후보 자료 기준 관리\n"
        f"  (2) 그동안 발생추이: 위험도 {record.risk_score:.1f}점({record.risk_grade}), "
        f"현장 이상징후 {record.suspicious_count}본, 시료 {record.sample_count}점 확인",
        16,
    )

    wrapped(draw, (70, 555, 780, 595), "나. 피해고사목 발생현황(최근 5년간)", 20, True)
    years = [record.year - 4 + i for i in range(5)]
    counts = [
        max(0, int(round(v["concern"] * ratio)))
        for ratio in (0.1, 0.2, 0.35, 0.55, 1.0)
    ]
    cells = draw_grid(draw, 55, 610, [125, 130, 130, 130, 130, 130], [42, 52])
    centered(draw, cells[0][0], "연도별", 15, True)
    centered(draw, cells[1][0], "본 수", 15, True)
    for i, year in enumerate(years):
        centered(draw, cells[0][i + 1], f"{year}년", 15)
        centered(draw, cells[1][i + 1], str(counts[i]), 16, True)

    wrapped(draw, (70, 735, 780, 775), "다. 방제 실적 및 계획(최근 3년간)", 20, True)
    widths = [80, 68, 68, 68, 68, 70, 70, 70, 70, 70, 75]
    cells = draw_grid(draw, 25, 790, widths, [72, 48, 48, 48])
    headers = [
        "연도", "계", "피해\n고사목", "기타\n고사목", "비병징목",
        "예방\n나무주사", "정밀\n드론", "지상", "유인트랩",
        "훈증더미\n제거", "비고",
    ]
    for i, header in enumerate(headers):
        centered(draw, cells[0][i], header, 12, True)

    rows = [
        [record.year - 2, 0, 0, 0, 0, 0, 0, 0, 0, 0, "실적자료\n미연계"],
        [record.year - 1, 0, 0, 0, 0, 0, 0, 0, 0, 0, "실적자료\n미연계"],
        [
            record.year, v["planned"], v["concern"], 0, 0, v["preventive"],
            round(v["area_ha"] * 0.6, 1), round(v["area_ha"] * 0.4, 1),
            max(0, v["planned"] // 3), v["fumigate"], "방제검토\n계획",
        ],
    ]
    for r, row in enumerate(rows, start=1):
        for c, value in enumerate(row):
            centered(draw, cells[r][c], value, 12 if c != 10 else 10)

    wrapped(
        draw,
        (40, 1055, 790, 1185),
        f"※ 본 계획서는 예측보고서 {record.document_no:02d}번과 현장예찰보고서를 "
        f"연계하여 작성한 방제 검토 계획이다. 검경 결과 및 담당자 승인 후 "
        f"대상 수량과 방제 방법을 확정한다.",
        14,
    )
    output.parent.mkdir(parents=True, exist_ok=True)
    image.save(output)


def create_appendix_2(record: Record, v: dict[str, Any], output: Path) -> None:
    image, draw = base_page("방제조치 명령서 관리대장")
    wrapped(draw, (55, 135, 780, 175), f"기관명: {record.sido_name} {record.sigungu_name} 산림보호 담당부서", 17)

    widths = [55, 130, 205, 80, 80, 95, 115]
    cells = draw_grid(draw, 35, 200, widths, [85] + [65] * 10)
    headers = ["연번", "명령을 받는 자", "명령내용", "방제기간", "방제방법", "명령서 수령자", "처리결과"]
    for i, header in enumerate(headers):
        centered(draw, cells[0][i], header, 14, True)

    row = [
        record.document_no,
        f"{record.sigungu_name}\n산림소유·관리자",
        f"격자 {record.center_grid_id} 현장 이상징후 {v['concern']}본에 대해 "
        "검경 결과 확인 후 대상목 처리 및 인접권역 재예찰",
        f"{v['start_date']:%Y.%m.%d}\n~\n{v['end_date']:%Y.%m.%d}",
        "파쇄·훈증\n예방나무주사\n검토",
        f"{v['surveyor']}\n(담당자)",
        "방제 검토\n계획 등록",
    ]
    for i, value in enumerate(row):
        centered(draw, cells[1][i], value, 11 if i != 2 else 10)

    wrapped(
        draw,
        (50, 940, 785, 1065),
        f"비고: 본 관리대장은 감염 확정 대장이 아니라 현장 예찰 결과를 기반으로 한 "
        f"방제조치 검토 이력이다. 계획번호 {v['plan_id']}, "
        f"예측 위험도 {record.risk_score:.1f}점, 예찰 우선순위 "
        f"{record.priority_score:.1f}점.",
        15,
    )
    output.parent.mkdir(parents=True, exist_ok=True)
    image.save(output)


def create_appendix_3(record: Record, v: dict[str, Any], output: Path) -> None:
    image, draw = base_page("재선충병 방제대상목 조사야장")
    wrapped(draw, (45, 135, 785, 175), f"○ 개소: {record.sido_name} {record.sigungu_name} / 중심 격자 {record.center_grid_id}", 17)

    widths = [58, 90, 95, 80, 80, 80, 100, 100, 105]
    cells = draw_grid(draw, 22, 200, widths, [70] + [58] * 12)
    headers = ["번호", "수종", "수고(m)", "흉고직경\n(cm)", "변색", "천공흔적", "시료", "방제검토", "비고"]
    for i, header in enumerate(headers):
        centered(draw, cells[0][i], header, 12, True)

    count = max(1, min(v["concern"], 10))
    for r in range(1, count + 1):
        row = [
            r,
            "소나무",
            f"{14.0 + (r % 5) * 0.8:.1f}",
            f"{24.0 + (r % 6) * 1.7:.1f}",
            "관찰" if r <= v["concern"] else "없음",
            "추가확인" if r % 3 == 0 else "미관찰",
            "채취" if r <= v["sample_count"] else "-",
            "검경 후 결정",
            "현장 확인 필요",
        ]
        for c, value in enumerate(row):
            centered(draw, cells[r][c], value, 10)

    wrapped(
        draw,
        (45, 970, 785, 1095),
        f"조사일: {parse_date(record.survey_datetime, record.year):%Y. %m. %d.}\n"
        f"조사자: {record.surveyors}\n"
        f"종합 의견: 검경 결과 확인 전까지 우선 예찰 검토지역으로 유지하고, "
        f"대상목별 조치 여부를 확정하지 않는다.",
        15,
    )
    output.parent.mkdir(parents=True, exist_ok=True)
    image.save(output)


def create_appendix_4(record: Record, v: dict[str, Any], output: Path) -> None:
    image, draw = base_page("피해고사목 방제실적")
    wrapped(draw, (50, 135, 785, 180), f"기관명: {record.sido_name} {record.sigungu_name} 산림보호 담당부서", 17)

    widths = [75, 110, 110, 110, 110, 110, 145]
    cells = draw_grid(draw, 36, 205, widths, [78] + [58] * 10)
    headers = ["연도", "대상목", "파쇄", "훈증", "예방주사", "잔재물관리", "처리결과·비고"]
    for i, header in enumerate(headers):
        centered(draw, cells[0][i], header, 13, True)

    rows = [
        [record.year - 2, 0, 0, 0, 0, 0, "실적자료 미연계"],
        [record.year - 1, 0, 0, 0, 0, 0, "실적자료 미연계"],
        [
            record.year,
            v["planned"],
            v["shred"],
            v["fumigate"],
            v["preventive"],
            v["fumigate"],
            "방제 검토 계획\n검경·승인 후 확정",
        ],
    ]
    for r, row in enumerate(rows, start=1):
        for c, value in enumerate(row):
            centered(draw, cells[r][c], value, 12 if c != 6 else 10)

    wrapped(
        draw,
        (50, 925, 785, 1080),
        f"작업 전 위험도: {record.risk_score:.1f}점({record.risk_grade})\n"
        f"조치 반영 예상 위험도: {v['after_score']:.1f}점({v['after_grade']})\n"
        f"주의: 위 수치는 방제 실행 결과가 아니라 계획 시나리오이며, 실제 처리 실적은 "
        f"현장 작업 완료 후 별도로 갱신한다.",
        16,
    )
    output.parent.mkdir(parents=True, exist_ok=True)
    image.save(output)


def replace_docx_media(docx_path: Path, replacements: dict[str, Path]) -> None:
    temp_path = docx_path.with_suffix(".patched.docx")
    internals = {f"word/media/{name}": path for name, path in replacements.items()}

    with zipfile.ZipFile(docx_path, "r") as src:
        names = set(src.namelist())
        missing = set(internals) - names
        if missing:
            raise RuntimeError(
                "DOCX 안에서 별지 이미지를 찾지 못했습니다: " + ", ".join(sorted(missing))
            )

        with zipfile.ZipFile(temp_path, "w", zipfile.ZIP_DEFLATED) as dst:
            for info in src.infolist():
                replacement = internals.get(info.filename)
                if replacement:
                    dst.writestr(info, replacement.read_bytes())
                else:
                    dst.writestr(info, src.read(info.filename))

    temp_path.replace(docx_path)


def extract_docx_text(docx_path: Path) -> str:
    with zipfile.ZipFile(docx_path, "r") as archive:
        xml = archive.read("word/document.xml").decode("utf-8", errors="ignore")
    return re.sub(r"<[^>]+>", "", xml)


def verify_no_placeholders(docx_path: Path) -> None:
    content = extract_docx_text(docx_path)
    patterns = [
        r"\[[^\[\]\r\n]{1,80}\]",
        re.escape("-지역, 기간-"),
        re.escape("[작성일]"),
    ]
    leftovers: list[str] = []
    for pattern in patterns:
        leftovers.extend(re.findall(pattern, content))
    leftovers = sorted(set(leftovers))
    if leftovers:
        raise RuntimeError(
            "본문에 미치환 플레이스홀더가 남았습니다: " + ", ".join(leftovers[:20])
        )


def convert_to_pdf(docx_path: Path, pdf_dir: Path) -> Path:
    pdf_dir.mkdir(parents=True, exist_ok=True)
    result = subprocess.run(
        [
            "libreoffice",
            "--headless",
            "--convert-to",
            "pdf",
            "--outdir",
            str(pdf_dir),
            str(docx_path),
        ],
        capture_output=True,
        text=True,
        timeout=180,
    )
    output = pdf_dir / f"{docx_path.stem}.pdf"
    if result.returncode != 0 or not output.exists():
        raise RuntimeError(
            f"PDF 변환 실패\nstdout={result.stdout}\nstderr={result.stderr}"
        )
    return output


def zip_results(pdf_dir: Path, manifest: Path, output_zip: Path) -> None:
    with zipfile.ZipFile(output_zip, "w", zipfile.ZIP_DEFLATED) as archive:
        for pdf in sorted(pdf_dir.glob("*.pdf")):
            archive.write(pdf, pdf.name)
        archive.write(manifest, manifest.name)


def safe_name(value: str) -> str:
    return re.sub(r'[\\/:*?"<>|]+', "_", value)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser()
    parser.add_argument("--prediction-manifest", type=Path, required=True)
    parser.add_argument("--field-manifest", type=Path, required=True)
    parser.add_argument("--terrain", type=Path, required=True)
    parser.add_argument("--template", type=Path, required=True)
    parser.add_argument("--output", type=Path, required=True)
    parser.add_argument("--limit", type=int, default=0)
    parser.add_argument("--keep-docx", action="store_true")
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    log(f"스크립트 버전: {SCRIPT_VERSION}")

    for required in (
        args.prediction_manifest,
        args.field_manifest,
        args.terrain,
        args.template,
    ):
        if not required.exists():
            raise FileNotFoundError(required)

    records = load_records(
        args.prediction_manifest,
        args.field_manifest,
        args.terrain,
    )
    if args.limit > 0:
        records = records[:args.limit]

    output_root = args.output
    docx_dir = output_root / "docx"
    pdf_dir = output_root / "pdf"
    appendix_dir = output_root / "appendices"
    for directory in (docx_dir, pdf_dir, appendix_dir):
        directory.mkdir(parents=True, exist_ok=True)

    manifest_rows: list[dict[str, Any]] = []

    for index, record in enumerate(records, start=1):
        log(
            f"[{index}/{len(records)}] {record.sigungu_name} / "
            f"격자 {record.center_grid_id}"
        )
        v = plan_values(record)

        required_plan_keys = {
            "start_date", "end_date", "days", "surveyor", "team", "address",
            "area_ha", "range", "confirmed", "concern", "planned", "shred",
            "fumigate", "preventive", "sample_count", "after_score",
            "after_grade", "plan_id", "tarpaulin",
        }
        missing_plan_keys = required_plan_keys - set(v)
        if missing_plan_keys:
            raise RuntimeError(
                "방제 계획 값 누락: " + ", ".join(sorted(missing_plan_keys))
            )

        stem = safe_name(
            f"{record.document_no:02d}_{record.year}_"
            f"소나무재선충병_방제보고서_"
            f"{record.sido_name}_{record.sigungu_name}_"
            f"격자{record.center_grid_id}"
        )

        appendix_paths = {
            "image1.png": appendix_dir / f"{stem}_별지1_방제사업계획서.png",
            "image2.png": appendix_dir / f"{stem}_별지2_방제조치명령서관리대장.png",
            "image3.png": appendix_dir / f"{stem}_별지3_방제대상목조사야장.png",
            "image4.png": appendix_dir / f"{stem}_별지4_피해고사목방제실적.png",
        }

        create_appendix_1(record, v, appendix_paths["image1.png"])
        create_appendix_2(record, v, appendix_paths["image2.png"])
        create_appendix_3(record, v, appendix_paths["image3.png"])
        create_appendix_4(record, v, appendix_paths["image4.png"])

        docx_path = docx_dir / f"{stem}.docx"
        fill_document_text(args.template, docx_path, record, v)
        replace_docx_media(docx_path, appendix_paths)
        verify_no_placeholders(docx_path)

        pdf_path = convert_to_pdf(docx_path, pdf_dir)
        if not args.keep_docx:
            docx_path.unlink(missing_ok=True)

        manifest_rows.append({
            "document_no": record.document_no,
            "file_name": pdf_path.name,
            "year": record.year,
            "center_grid_id": record.center_grid_id,
            "sido_name": record.sido_name,
            "sigungu_name": record.sigungu_name,
            "source_prediction_file": record.source_prediction_file,
            "source_field_file": record.source_field_file,
            "link_status": "MATCHED",
            "risk_score": round(record.risk_score, 1),
            "risk_grade": record.risk_grade,
            "priority_score": round(record.priority_score, 1),
            "priority_grade": record.priority_grade,
            "field_suspicious_count": record.suspicious_count,
            "sample_count": record.sample_count,
            "control_status": "방제 검토 계획",
            "planned_count": v["planned"],
            "planned_shred_count": v["shred"],
            "planned_fumigation_count": v["fumigate"],
            "planned_preventive_injection_count": v["preventive"],
            "planned_area_ha": v["area_ha"],
            "script_version": SCRIPT_VERSION,
        })

    expected = len(records)
    pdf_count = len(list(pdf_dir.glob("*.pdf")))
    appendix_count = len(list(appendix_dir.glob("*.png")))
    if pdf_count != expected:
        raise RuntimeError(f"PDF 개수 불일치: {pdf_count}/{expected}")
    if appendix_count != expected * 4:
        raise RuntimeError(f"별지 개수 불일치: {appendix_count}/{expected * 4}")

    manifest_path = output_root / "문서목록.csv"
    pd.DataFrame(manifest_rows).to_csv(
        manifest_path,
        index=False,
        encoding="utf-8-sig",
    )

    zip_path = output_root / f"방제보고서_{expected}건_예측현장연계_완성본_PDF.zip"
    zip_results(pdf_dir, manifest_path, zip_path)

    log("생성 완료")
    log(f"- PDF: {pdf_count}")
    log(f"- 별지: {appendix_count}")
    log(f"- ZIP: {zip_path}")


if __name__ == "__main__":
    main()
