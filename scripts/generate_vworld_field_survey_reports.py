#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
기존 「소나무재선충 현장 예찰 보고서」 DOCX 양식을 유지하면서
VWorld 배경지도, 중심 격자 + 주변 8개 격자(3x3), 현장 이동경로,
관찰지점, 시료채취지점을 삽입해 PDF 30건을 생성한다.

전제
- scripts/generate_vworld_prediction_reports.py가 같은 scripts 폴더에 존재한다.
- /opt/pine-wilt/rag-backend/.env에 VWORLD_API_KEY, VWORLD_API_DOMAIN이 있다.
- LibreOffice가 설치되어 있다.

실행 예시
/opt/pine-wilt/report-venv/bin/python \
  scripts/generate_vworld_field_survey_reports.py \
  --terrain data/terrain_pine_site_features_south_500m.csv \
  --infection public/data/infection_history_2016_2021.geojson \
  --sigungu data/sigungu_boundary.geojson \
  --template 'rag-backend/data/report_templates/[양식]소나무재선충병 현장 예찰 보고서_빈양식.docx' \
  --prediction-manifest rag-backend/data/generated_reports/prediction_30/문서목록.csv \
  --output rag-backend/data/generated_reports/field_survey_30 \
  --zoom 10
"""

from __future__ import annotations

import argparse
import importlib.util
import io
import math
import os
import random
import re
import shutil
import subprocess
import sys
import time
import zipfile
from dataclasses import asdict, dataclass
from pathlib import Path
from typing import Any

import geopandas as gpd
import numpy as np
import pandas as pd
import requests
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
from docx.shared import Mm
from dotenv import load_dotenv
from PIL import Image, ImageDraw, ImageFont
from pyproj import Transformer
from shapely.geometry import Point, Polygon


# -----------------------------------------------------------------------------
# 기존 발생 예측 보고서 스크립트의 검증된 공통 로직 재사용
# -----------------------------------------------------------------------------

def load_prediction_module() -> Any:
    module_path = Path(__file__).with_name("generate_vworld_prediction_reports.py")
    if not module_path.exists():
        raise FileNotFoundError(
            f"공통 모듈을 찾지 못했습니다: {module_path}\n"
            "scripts/generate_vworld_prediction_reports.py가 같은 폴더에 있어야 합니다."
        )
    spec = importlib.util.spec_from_file_location("prediction_report_common", module_path)
    if spec is None or spec.loader is None:
        raise RuntimeError(f"공통 모듈 로딩 실패: {module_path}")
    module = importlib.util.module_from_spec(spec)
    sys.modules[spec.name] = module
    spec.loader.exec_module(module)
    return module


COMMON = load_prediction_module()

YEAR_RANGE = COMMON.YEAR_RANGE
GRID_SIZE_M = COMMON.GRID_SIZE_M
MAP_WIDTH = 1024
MAP_HEIGHT = 704
DEFAULT_ZOOM = 10
SCRIPT_VERSION = "2026.07.20-linked-field-appendices-v2"

TerrainCell = COMMON.TerrainCell
ReportRecord = COMMON.ReportRecord


# -----------------------------------------------------------------------------
# 현장 예찰 보고서용 데이터 모델
# -----------------------------------------------------------------------------

@dataclass
class FieldSurveyData:
    survey_datetime: str
    weather: str
    temperature_c: float
    wind_speed_ms: float
    organization_team: str
    surveyors: str
    discovery_route: str
    address: str
    forest_compartment: str
    latitude: float
    longitude: float
    ai_result: str
    overall_judgment: str
    species: str
    total_trees: int
    detail_classification: str
    tree_height_m: float
    dbh_cm: float
    discoloration_stage: str
    vector_trace: str
    bark_wood_observation: str
    investigator_opinion: str
    sample_description: str
    qr_code: str
    system_link_result: str
    followup_date: str
    inspection_agency_status: str
    followup_plan: str
    route_type: str
    sample_count: int
    suspicious_count: int


# -----------------------------------------------------------------------------
# 공통 유틸
# -----------------------------------------------------------------------------

def log(message: str) -> None:
    print(message, flush=True)


def sanitize_filename(value: str) -> str:
    return re.sub(r'[\\/:*?"<>|]+', "_", value).strip()


def load_linked_prediction_records(
    manifest_path: Path,
) -> tuple[list[ReportRecord], dict[int, dict[str, Any]]]:
    """발생 예측 보고서 문서목록을 읽어 동일한 30개 격자를 같은 순서로 반환한다."""
    log(f"예측보고서 연계 목록 로딩: {manifest_path}")
    frame = pd.read_csv(manifest_path, encoding="utf-8-sig")

    required = {
        "document_no",
        "file_name",
        "year",
        "center_grid_id",
        "sido_name",
        "sigungu_name",
        "center_annual_count",
        "center_cumulative_count",
        "risk_score",
        "risk_grade",
        "priority_score",
        "priority_grade",
    }
    missing = required - set(frame.columns)
    if missing:
        raise ValueError(
            "예측보고서 문서목록.csv 필수 컬럼이 없습니다: " + ", ".join(sorted(missing))
        )

    if frame.empty:
        raise ValueError("예측보고서 문서목록.csv가 비어 있습니다.")

    frame = frame.sort_values("document_no").reset_index(drop=True)
    if frame["document_no"].duplicated().any():
        raise ValueError("예측보고서 document_no가 중복되었습니다.")
    if frame["center_grid_id"].duplicated().any():
        duplicates = frame.loc[frame["center_grid_id"].duplicated(), "center_grid_id"].tolist()
        raise ValueError(f"예측보고서 중심 격자가 중복되었습니다: {duplicates}")

    records: list[ReportRecord] = []
    metadata_by_no: dict[int, dict[str, Any]] = {}
    for row in frame.to_dict(orient="records"):
        document_no = int(row["document_no"])
        record = ReportRecord(
            report_no=document_no,
            year=int(row["year"]),
            center_grid_id=int(row["center_grid_id"]),
            annual_count=int(row["center_annual_count"]),
            cumulative_count=int(row["center_cumulative_count"]),
            sido_name=str(row["sido_name"]),
            sigungu_name=str(row["sigungu_name"]),
        )
        records.append(record)
        metadata_by_no[document_no] = row

    log(f"예측보고서 {len(records)}건과 1:1 연계 완료")
    return records, metadata_by_no


def apply_prediction_manifest_metrics(
    metrics: dict[str, Any],
    prediction_row: dict[str, Any],
) -> dict[str, Any]:
    """예측보고서에 기록된 위험도·우선순위를 현장 예찰 보고서에 그대로 승계한다."""
    linked = dict(metrics)
    linked["risk_score"] = float(prediction_row["risk_score"])
    linked["risk_grade"] = str(prediction_row["risk_grade"])
    linked["priority_score"] = float(prediction_row["priority_score"])
    linked["priority_grade"] = str(prediction_row["priority_grade"])

    manifest_grids = str(prediction_row.get("block_grid_ids", "")).strip()
    if manifest_grids:
        manifest_ids = [int(value) for value in manifest_grids.split("|") if str(value).strip()]
        calculated_ids = [int(value) for value in linked["block_grid_ids"]]
        if set(manifest_ids) != set(calculated_ids):
            raise RuntimeError(
                "예측보고서와 현장예찰의 3x3 격자 구성이 다릅니다. "
                f"예측={manifest_ids}, 현장계산={calculated_ids}"
            )
        linked["block_grid_ids"] = manifest_ids
        linked["block_count"] = len(manifest_ids)
    return linked


def find_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    candidates = [
        "/usr/share/fonts/truetype/nanum/NanumGothicBold.ttf" if bold else "/usr/share/fonts/truetype/nanum/NanumGothic.ttf",
        "/usr/share/fonts/truetype/nanum/NanumSquareRoundB.ttf" if bold else "/usr/share/fonts/truetype/nanum/NanumSquareRoundR.ttf",
        "/usr/share/fonts/truetype/unfonts-core/UnDotumBold.ttf" if bold else "/usr/share/fonts/truetype/unfonts-core/UnDotum.ttf",
    ]
    for path in candidates:
        if Path(path).exists():
            return ImageFont.truetype(path, size=size)
    return ImageFont.load_default()


def deterministic_rng(record: ReportRecord) -> random.Random:
    return random.Random(record.center_grid_id * 10000 + record.year)


def add_days_text(year: int, month: int, day: int, offset: int) -> str:
    # 12월 날짜만 사용하며 28일을 넘지 않도록 순환
    adjusted = ((day - 1 + offset) % 28) + 1
    return f"{year}. {month:02d}. {adjusted:02d}."


# -----------------------------------------------------------------------------
# 현장 관찰값 생성 규칙
# -----------------------------------------------------------------------------

def create_field_survey_data(
    record: ReportRecord,
    metrics: dict[str, Any],
    center_cell: TerrainCell,
) -> FieldSurveyData:
    rng = deterministic_rng(record)
    transformer = Transformer.from_crs("EPSG:5186", "EPSG:4326", always_xy=True)
    center = center_cell.geometry.centroid
    lon, lat = transformer.transform(center.x, center.y)

    risk = float(metrics["risk_score"])
    pressure = float(metrics["infection_pressure"])
    pine = float(metrics["pine_mean"])

    # 위험도별 관찰 강도
    if risk >= 85:
        suspicious = rng.randint(5, 9)
        route_type = "현장 정밀예찰 및 드론 보조 촬영"
        discoloration = rng.choice([
            "수관 상부 황화와 부분 갈변이 함께 관찰되는 중기 단계",
            "가지 끝부터 갈변이 확산된 중기~후기 단계",
        ])
        vector_trace = rng.choice([
            "일부 수간에서 소형 천공 흔적이 관찰되어 추가 확인 필요",
            "수피 표면에 천공 의심 흔적이 확인되어 정밀 관찰 실시",
        ])
        sample_count = rng.randint(3, 5)
    elif risk >= 70:
        suspicious = rng.randint(2, 6)
        route_type = "드론 선행예찰 후 현장 확인"
        discoloration = rng.choice([
            "수관 일부 황화가 확인되는 초기~중기 단계",
            "부분 갈변 및 잎 처짐이 관찰되는 초기 단계",
        ])
        vector_trace = rng.choice([
            "뚜렷한 탈출공은 미확인되었으나 일부 천공 의심 흔적 관찰",
            "수피 틈과 가지 분지부를 중심으로 추가 확인 필요",
        ])
        sample_count = rng.randint(2, 4)
    elif risk >= 55:
        suspicious = rng.randint(1, 3)
        route_type = "드론 우선예찰 및 표본 현장점검"
        discoloration = rng.choice([
            "경미한 잎 변색이 확인되는 초기 단계",
            "부분적인 황화가 있으나 계절성 변화 여부 추가 확인 필요",
        ])
        vector_trace = "뚜렷한 매개충 흔적은 확인되지 않음"
        sample_count = rng.randint(1, 2)
    else:
        suspicious = rng.randint(0, 2)
        route_type = "정기 순찰 및 드론 모니터링"
        discoloration = "대부분 정상이며 일부 개체에서 경미한 변색 관찰"
        vector_trace = "매개충 흔적 미확인"
        sample_count = 1 if suspicious > 0 else 0

    normal_count = rng.randint(8, 18)
    total = suspicious + normal_count
    dead_count = max(0, suspicious - rng.randint(0, 2))
    yellow_count = suspicious - dead_count

    species = rng.choice(["소나무", "곰솔", "잣나무", "리기다소나무"])
    height = round(rng.uniform(8.0, 17.5), 1)
    dbh = round(rng.uniform(18.0, 37.0), 1)
    temperature = round(rng.uniform(4.0, 17.0), 1)
    wind = round(rng.uniform(0.8, 4.8), 1)
    weather = rng.choice(["맑음", "구름 조금", "흐림", "맑은 후 구름 많음"])

    team_number = (record.report_no % 4) + 1
    team_names = [
        ("산림보호 예찰 1조", "김도윤, 이서진"),
        ("산림보호 예찰 2조", "박지훈, 최유나"),
        ("산림병해충 대응 1조", "정민수, 한지우"),
        ("현장예찰 지원 2조", "강서준, 윤가은"),
    ]
    organization, surveyors = team_names[record.report_no % len(team_names)]

    route = (
        f"AI 우선 예찰 검토지역 지정 후 {route_type} 방식으로 중심 격자와 주변 8개 격자를 순차 확인"
    )
    compartment = f"{(record.center_grid_id % 90) + 10}임반-{(record.center_grid_id % 9) + 1}소반"

    if suspicious > 0:
        bark_obs = rng.choice([
            "의심 개체의 수피 일부 건조와 목질부 수분 저하가 관찰됨",
            "수피 박리부 주변에서 갈변이 확인되어 시료 채취 실시",
            "가지 절단면 일부에서 변색이 관찰되어 목편 시료 확보",
        ])
        opinion = (
            f"중심 격자를 포함한 3×3 권역에서 변색 또는 고사 의심 개체 {suspicious}본을 확인했습니다. "
            f"현장만으로 감염 여부를 판단하기 어려워 시료를 채취했으며, 인접 격자까지 추가 모니터링이 필요합니다."
        )
    else:
        bark_obs = "수피 및 목질부에서 특이사항이 확인되지 않음"
        opinion = (
            "현재 조사 범위에서는 뚜렷한 감염 의심 징후가 확인되지 않았습니다. "
            "다만 과거 발생 이력이 있는 권역이므로 정기 모니터링을 유지하겠습니다."
        )

    qr_code = f"FS-{record.year}-{record.center_grid_id}-{record.report_no:02d}"
    if sample_count > 0:
        sample_desc = f"목편 {sample_count}점, 가지 시료 {max(1, sample_count - 1)}점"
        system_result = f"현장 예찰 시스템에 사진·좌표·시료정보 등록 완료, QR코드 {qr_code} 연동"
        agency_status = "관할 산림환경연구기관 검경 의뢰 접수, 결과 대기"
        followup_plan = (
            "검경 결과 확인 전까지 대상 3×3 권역을 우선 예찰 검토지역으로 유지하고, "
            "의심 개체 주변 반경을 재확인한다. 검경 결과에 따라 방제 검토와 인접 권역 확대조사를 연계한다."
        )
    else:
        sample_desc = "현장 관찰 결과 시료 미채취"
        system_result = f"현장 사진·좌표·조사결과 등록 완료, 조사기록 {qr_code} 생성"
        agency_status = "검경 의뢰 없음, 정기 예찰 대상으로 관리"
        followup_plan = (
            "2주 이내 동일 권역을 재확인하고 변색 진행 여부를 비교한다. "
            "새로운 의심 징후 확인 시 즉시 시료 채취와 검경 의뢰를 시행한다."
        )

    survey_day = 8 + (record.report_no % 18)
    survey_datetime = f"{record.year}. 12. {survey_day:02d}. 09:30~14:30"
    followup_date = add_days_text(record.year, 12, survey_day, 7)

    detail = f"정상 관찰 {normal_count}본, 변색 의심 {yellow_count}본, 고사 의심 {dead_count}본"
    ai_result = (
        f"종합 위험도 {risk:.1f}점({metrics['risk_grade']}), 감염압력 {pressure:.1f}점, "
        f"예찰 우선순위 {metrics['priority_score']:.1f}점({metrics['priority_grade']})"
    )
    judgment = (
        f"과거 발생 이력과 3×3 권역 내 위험 신호를 고려할 때 현장 확인이 필요함. "
        f"권역 평균 소나무류 비율은 {pine:.1f}%이며, 중심 격자와 인접 격자의 연속 예찰을 권고함."
    )

    return FieldSurveyData(
        survey_datetime=survey_datetime,
        weather=weather,
        temperature_c=temperature,
        wind_speed_ms=wind,
        organization_team=organization,
        surveyors=surveyors,
        discovery_route=route,
        address=f"{record.sido_name} {record.sigungu_name} 산림 일원",
        forest_compartment=compartment,
        latitude=lat,
        longitude=lon,
        ai_result=ai_result,
        overall_judgment=judgment,
        species=species,
        total_trees=total,
        detail_classification=detail,
        tree_height_m=height,
        dbh_cm=dbh,
        discoloration_stage=discoloration,
        vector_trace=vector_trace,
        bark_wood_observation=bark_obs,
        investigator_opinion=opinion,
        sample_description=sample_desc,
        qr_code=qr_code,
        system_link_result=system_result,
        followup_date=followup_date,
        inspection_agency_status=agency_status,
        followup_plan=followup_plan,
        route_type=route_type,
        sample_count=sample_count,
        suspicious_count=suspicious,
    )


# -----------------------------------------------------------------------------
# VWorld 지도 + 현장 이동경로/관찰점/시료점 오버레이
# -----------------------------------------------------------------------------

def projected_point_to_image(
    point_5186: Point,
    center_lon: float,
    center_lat: float,
    zoom: int,
    transformer: Transformer,
) -> tuple[float, float]:
    lon, lat = transformer.transform(point_5186.x, point_5186.y)
    center_x, center_y = COMMON.lonlat_to_world_pixel(center_lon, center_lat, zoom)
    wx, wy = COMMON.lonlat_to_world_pixel(lon, lat, zoom)
    return MAP_WIDTH / 2 + (wx - center_x), MAP_HEIGHT / 2 + (wy - center_y)


def make_survey_points(
    record: ReportRecord,
    center_cell: TerrainCell,
    survey: FieldSurveyData,
) -> tuple[list[Point], list[Point], list[Point]]:
    rng = deterministic_rng(record)
    center = center_cell.geometry.centroid

    # 3x3 권역 내부의 이동경로 6개 점
    route_points = [
        Point(center.x - 600, center.y - 500),
        Point(center.x - 350, center.y + 450),
        Point(center.x + 100, center.y + 550),
        Point(center.x + 550, center.y + 250),
        Point(center.x + 500, center.y - 400),
        Point(center.x, center.y - 550),
    ]

    observation_count = max(1, min(5, survey.suspicious_count))
    observation_points: list[Point] = []
    for _ in range(observation_count):
        observation_points.append(
            Point(center.x + rng.uniform(-620, 620), center.y + rng.uniform(-620, 620))
        )

    sample_points = observation_points[: min(survey.sample_count, len(observation_points))]
    return route_points, observation_points, sample_points


def build_field_survey_map(
    output_path: Path,
    record: ReportRecord,
    cells: list[TerrainCell],
    infection: gpd.GeoDataFrame,
    survey: FieldSurveyData,
    api_key: str,
    domain: str,
    zoom: int,
    basemap: str,
) -> None:
    transformer = Transformer.from_crs("EPSG:5186", "EPSG:4326", always_xy=True)
    center_cell = next(cell for cell in cells if cell.grid_id == record.center_grid_id)
    center_point = center_cell.geometry.centroid
    center_lon, center_lat = transformer.transform(center_point.x, center_point.y)

    image = COMMON.request_vworld_map(
        api_key=api_key,
        domain=domain,
        center_lon=center_lon,
        center_lat=center_lat,
        zoom=zoom,
        basemap=basemap,
        width=MAP_WIDTH,
        height=MAP_HEIGHT,
    )
    overlay = Image.new("RGBA", image.size, (0, 0, 0, 0))
    draw = ImageDraw.Draw(overlay, "RGBA")

    # 주변 발생 이력 오버레이
    search_area = center_point.buffer(14_000)
    local = infection[infection.intersects(search_area)].copy()
    annual_col = f"infection_count_{record.year}"
    for _, row in local.iterrows():
        annual = int(row.get(annual_col, 0))
        cumulative = int(row.get("infection_count_2016_2021", 0))
        if annual <= 0 and cumulative <= 0:
            continue
        grade = COMMON.local_risk_grade(annual, cumulative)
        points = COMMON.geometry_to_image_points(
            row.geometry,
            center_lon,
            center_lat,
            zoom,
            MAP_WIDTH,
            MAP_HEIGHT,
            transformer,
        )
        if points and any(-100 <= x <= MAP_WIDTH + 100 and -100 <= y <= MAP_HEIGHT + 100 for x, y in points):
            color = COMMON.RISK_COLORS[grade]
            draw.polygon(points, fill=(color[0], color[1], color[2], 120), outline=(255, 255, 255, 90), width=1)

    # 3x3 격자 강조
    for cell in cells:
        points = COMMON.geometry_to_image_points(
            cell.geometry,
            center_lon,
            center_lat,
            zoom,
            MAP_WIDTH,
            MAP_HEIGHT,
            transformer,
        )
        if cell.grid_id == record.center_grid_id:
            draw.polygon(points, fill=(255, 65, 65, 75), outline=(170, 0, 0, 255), width=5)
        else:
            draw.line(points, fill=(20, 82, 190, 240), width=4, joint="curve")

    block_union = gpd.GeoSeries([cell.geometry for cell in cells], crs="EPSG:5186").union_all()
    if block_union.geom_type == "Polygon":
        outer = COMMON.geometry_to_image_points(
            block_union,
            center_lon,
            center_lat,
            zoom,
            MAP_WIDTH,
            MAP_HEIGHT,
            transformer,
        )
        draw.line(outer, fill=(0, 45, 130, 255), width=6, joint="curve")

    # 이동경로 및 관찰점
    route_points, observation_points, sample_points = make_survey_points(record, center_cell, survey)
    route_pixels = [projected_point_to_image(p, center_lon, center_lat, zoom, transformer) for p in route_points]
    if len(route_pixels) >= 2:
        draw.line(route_pixels, fill=(20, 20, 20, 230), width=5, joint="curve")
        for i in range(len(route_pixels) - 1):
            x1, y1 = route_pixels[i]
            x2, y2 = route_pixels[i + 1]
            # 진행 방향 표시
            mx = x1 + (x2 - x1) * 0.72
            my = y1 + (y2 - y1) * 0.72
            ang = math.atan2(y2 - y1, x2 - x1)
            size = 12
            p1 = (mx, my)
            p2 = (mx - size * math.cos(ang - 0.5), my - size * math.sin(ang - 0.5))
            p3 = (mx - size * math.cos(ang + 0.5), my - size * math.sin(ang + 0.5))
            draw.polygon([p1, p2, p3], fill=(20, 20, 20, 230))

    for point in observation_points:
        x, y = projected_point_to_image(point, center_lon, center_lat, zoom, transformer)
        draw.polygon([(x, y - 13), (x - 11, y + 9), (x + 11, y + 9)], fill=(255, 45, 45, 245), outline=(120, 0, 0, 255))

    for point in sample_points:
        x, y = projected_point_to_image(point, center_lon, center_lat, zoom, transformer)
        draw.rectangle((x - 8, y - 8, x + 8, y + 8), fill=(35, 85, 220, 255), outline=(255, 255, 255, 255), width=2)

    # 시작점
    sx, sy = route_pixels[0]
    draw.ellipse((sx - 10, sy - 10, sx + 10, sy + 10), fill=(35, 190, 95, 255), outline=(255, 255, 255, 255), width=2)

    title_font = find_font(30, bold=True)
    label_font = find_font(22, bold=True)
    small_font = find_font(17, bold=False)
    title = f"{record.sido_name} {record.sigungu_name} / 현장 예찰 조사도"
    bbox = draw.textbbox((0, 0), title, font=title_font)
    tw = bbox[2] - bbox[0]
    draw.rounded_rectangle(
        (MAP_WIDTH / 2 - tw / 2 - 16, 14, MAP_WIDTH / 2 + tw / 2 + 16, 60),
        radius=9,
        fill=(255, 255, 255, 230),
        outline=(150, 150, 150, 220),
        width=2,
    )
    draw.text((MAP_WIDTH / 2 - tw / 2, 20), title, font=title_font, fill=(20, 20, 20, 255))

    center_label = f"중심 격자 {record.center_grid_id}"
    draw.rounded_rectangle((MAP_WIDTH / 2 + 12, MAP_HEIGHT / 2 - 52, MAP_WIDTH / 2 + 250, MAP_HEIGHT / 2 - 14), radius=7, fill=(255, 255, 255, 235), outline=(160, 160, 160, 255), width=2)
    draw.text((MAP_WIDTH / 2 + 20, MAP_HEIGHT / 2 - 48), center_label, font=label_font, fill=(20, 20, 20, 255))

    # 범례
    lx, ly, lw, lh = 20, MAP_HEIGHT - 158, 560, 135
    draw.rounded_rectangle((lx, ly, lx + lw, ly + lh), radius=10, fill=(255, 255, 255, 238), outline=(165, 165, 165, 230), width=2)
    legend = [
        ("예찰 시작점", "circle", (35, 190, 95, 255)),
        ("감염 의심목 관찰점", "triangle", (255, 45, 45, 245)),
        ("시료 채취점", "square", (35, 85, 220, 255)),
        ("현장 이동경로", "line", (20, 20, 20, 230)),
        ("중심 격자", "box_red", (255, 65, 65, 75)),
        ("주변 8개 격자", "box_blue", (20, 82, 190, 240)),
    ]
    for idx, (name, shape, color) in enumerate(legend):
        col = idx % 2
        row = idx // 2
        x = lx + 18 + col * 275
        y = ly + 17 + row * 38
        if shape == "circle":
            draw.ellipse((x, y, x + 22, y + 22), fill=color)
        elif shape == "triangle":
            draw.polygon([(x + 11, y), (x, y + 22), (x + 22, y + 22)], fill=color)
        elif shape == "square":
            draw.rectangle((x, y, x + 22, y + 22), fill=color)
        elif shape == "line":
            draw.line((x, y + 11, x + 26, y + 11), fill=color, width=4)
        elif shape == "box_red":
            draw.rectangle((x, y, x + 24, y + 22), fill=(255, 65, 65, 75), outline=(170, 0, 0, 255), width=3)
        else:
            draw.rectangle((x, y, x + 24, y + 22), fill=(255, 255, 255, 20), outline=(20, 82, 190, 240), width=3)
        draw.text((x + 34, y - 1), name, font=small_font, fill=(20, 20, 20, 255))

    source = "배경지도: VWorld | 예찰권역: 중심 격자 포함 3×3"
    sb = draw.textbbox((0, 0), source, font=small_font)
    sw = sb[2] - sb[0]
    draw.rounded_rectangle((MAP_WIDTH - sw - 28, MAP_HEIGHT - 34, MAP_WIDTH - 10, MAP_HEIGHT - 6), radius=5, fill=(255, 255, 255, 225))
    draw.text((MAP_WIDTH - sw - 20, MAP_HEIGHT - 32), source, font=small_font, fill=(50, 50, 50, 255))

    result = Image.alpha_composite(image, overlay).convert("RGB")
    output_path.parent.mkdir(parents=True, exist_ok=True)
    result.save(output_path, quality=95)



# -----------------------------------------------------------------------------
# 7~8페이지 별지 이미지 생성 및 DOCX 내부 이미지 교체
# -----------------------------------------------------------------------------

def draw_text_fit(
    draw: ImageDraw.ImageDraw,
    box: tuple[int, int, int, int],
    text: str,
    *,
    max_size: int = 22,
    min_size: int = 11,
    bold: bool = False,
    align: str = "center",
    line_spacing: int = 3,
    fill: tuple[int, int, int, int] | tuple[int, int, int] = (20, 20, 20),
) -> None:
    """지정 셀 안에서 자동 줄바꿈/축소해 텍스트를 그린다."""
    x1, y1, x2, y2 = box
    width = max(1, x2 - x1 - 8)
    height = max(1, y2 - y1 - 6)
    raw = str(text).strip()
    if not raw:
        return

    def wrap_text(font: ImageFont.ImageFont) -> list[str]:
        lines: list[str] = []
        for source_line in raw.splitlines() or [""]:
            if not source_line:
                lines.append("")
                continue
            current = ""
            for ch in source_line:
                trial = current + ch
                bbox = draw.textbbox((0, 0), trial, font=font)
                if bbox[2] - bbox[0] <= width or not current:
                    current = trial
                else:
                    lines.append(current)
                    current = ch
            if current:
                lines.append(current)
        return lines or [""]

    chosen_font = find_font(min_size, bold=bold)
    chosen_lines = [raw]
    for size in range(max_size, min_size - 1, -1):
        font = find_font(size, bold=bold)
        lines = wrap_text(font)
        line_heights = []
        for line in lines:
            bbox = draw.textbbox((0, 0), line or "가", font=font)
            line_heights.append(max(1, bbox[3] - bbox[1]))
        total_h = sum(line_heights) + line_spacing * max(0, len(lines) - 1)
        if total_h <= height:
            chosen_font = font
            chosen_lines = lines
            break

    line_metrics = []
    for line in chosen_lines:
        bbox = draw.textbbox((0, 0), line or "가", font=chosen_font)
        line_metrics.append((bbox[2] - bbox[0], max(1, bbox[3] - bbox[1])))
    total_h = sum(h for _, h in line_metrics) + line_spacing * max(0, len(chosen_lines) - 1)
    y = y1 + max(0, (height - total_h) / 2) + 2

    for line, (line_w, line_h) in zip(chosen_lines, line_metrics):
        if align == "left":
            x = x1 + 5
        elif align == "right":
            x = x2 - line_w - 5
        else:
            x = x1 + max(0, ((x2 - x1) - line_w) / 2)
        draw.text((x, y), line, font=chosen_font, fill=fill)
        y += line_h + line_spacing


def extract_template_media(template_path: Path, media_name: str) -> Image.Image:
    """DOCX 템플릿의 word/media 이미지를 읽는다."""
    internal = f"word/media/{media_name}"
    with zipfile.ZipFile(template_path, "r") as archive:
        if internal not in archive.namelist():
            raise RuntimeError(f"템플릿에서 별지 이미지를 찾지 못했습니다: {internal}")
        raw = archive.read(internal)
    return Image.open(io.BytesIO(raw)).convert("RGB")


def build_air_survey_plan_appendix(
    template_path: Path,
    output_path: Path,
    record: ReportRecord,
    survey: FieldSurveyData,
    metrics: dict[str, Any],
) -> None:
    """7페이지: 유인항공예찰 계획 별지 작성."""
    image = extract_template_media(template_path, "image1.png")
    if image.size != (826, 1264):
        raise RuntimeError(
            f"유인항공예찰 계획 이미지 크기가 예상과 다릅니다: {image.size}, 예상=(826, 1264)"
        )
    draw = ImageDraw.Draw(image)

    # 첫 번째 데이터 행 좌표(원본 826×1264 기준)
    boxes = {
        "sigungu": (24, 239, 130, 286),
        "region": (130, 239, 210, 286),
        "area": (210, 239, 295, 286),
        "date": (295, 239, 374, 286),
        "landing": (374, 239, 470, 286),
        "org": (470, 239, 553, 286),
        "rank": (553, 239, 637, 286),
        "name": (637, 239, 722, 286),
        "helicopter": (722, 239, 814, 286),
        "attachment": (24, 1115, 814, 1193),
    }

    # 원본 양식에 들어 있던 파란 안내문 및 기본 문구를 지운 뒤 셀 내부만 다시 작성
    for box in boxes.values():
        x1, y1, x2, y2 = box
        draw.rectangle((x1 + 2, y1 + 2, x2 - 2, y2 - 2), fill=(255, 255, 255))

    survey_date = survey.survey_datetime.split(". 09:")[0].strip()
    area_ha = 9 * (GRID_SIZE_M * GRID_SIZE_M) / 10_000
    primary_name = survey.surveyors.split(",")[0].strip()
    org_name = "산림보호과"
    risk_note = (
        f"중심 격자 {record.center_grid_id}\n"
        f"주변 8개 포함"
    )

    draw_text_fit(draw, boxes["sigungu"], record.sigungu_name, max_size=18, min_size=12, bold=True)
    draw_text_fit(draw, boxes["region"], risk_note, max_size=15, min_size=10)
    draw_text_fit(draw, boxes["area"], f"{area_ha:.0f}", max_size=18, min_size=12)
    draw_text_fit(draw, boxes["date"], survey_date, max_size=15, min_size=10)
    draw_text_fit(draw, boxes["landing"], f"{record.sigungu_name}\n임시착륙장", max_size=14, min_size=10)
    draw_text_fit(draw, boxes["org"], org_name, max_size=14, min_size=10)
    draw_text_fit(draw, boxes["rank"], "주무관", max_size=16, min_size=11)
    draw_text_fit(draw, boxes["name"], primary_name, max_size=16, min_size=11)
    draw_text_fit(draw, boxes["helicopter"], "산림청\n중형헬기", max_size=14, min_size=10)
    draw_text_fit(
        draw,
        boxes["attachment"],
        (
            f"중심 격자 {record.center_grid_id} 포함 3×3 예찰권역, "
            f"계획면적 {area_ha:.0f}ha. 본문 현장 예찰 조사도 및 이동경로 참조."
        ),
        max_size=17,
        min_size=12,
        align="left",
    )

    output_path.parent.mkdir(parents=True, exist_ok=True)
    image.save(output_path, format="PNG")


def build_air_survey_result_appendix(
    template_path: Path,
    output_path: Path,
    record: ReportRecord,
    survey: FieldSurveyData,
    metrics: dict[str, Any],
) -> None:
    """8페이지: 유인항공예찰 조사결과 별지 작성."""
    image = extract_template_media(template_path, "image2.png")
    if image.size != (832, 1228):
        raise RuntimeError(
            f"유인항공예찰 조사결과 이미지 크기가 예상과 다릅니다: {image.size}, 예상=(832, 1228)"
        )
    draw = ImageDraw.Draw(image)

    # 상단 기본정보
    # 상단 빈칸에 남아 있는 원본 예시문을 먼저 제거
    draw.rectangle((145, 143, 810, 190), fill=(255, 255, 255))
    draw.rectangle((150, 198, 810, 245), fill=(255, 255, 255))
    draw.rectangle((300, 250, 810, 296), fill=(255, 255, 255))

    draw_text_fit(
        draw,
        (145, 145, 785, 188),
        f"{record.sido_name} {record.sigungu_name} 산림보호 담당부서",
        max_size=21,
        min_size=14,
        align="left",
        bold=True,
    )
    survey_date = survey.survey_datetime.split(" 09:")[0].strip()
    draw_text_fit(
        draw,
        (150, 201, 790, 242),
        f"{survey_date} ~ {survey_date} (1일)",
        max_size=20,
        min_size=13,
        align="left",
    )
    draw_text_fit(
        draw,
        (300, 253, 790, 293),
        "1대(산림청 중형헬기)",
        max_size=20,
        min_size=13,
        align="left",
    )

    # 첫 번째 결과 행
    boxes = {
        "sigungu": (20, 507, 121, 584),
        "region": (121, 507, 296, 584),
        "area": (296, 507, 402, 584),
        "total": (402, 507, 495, 584),
        "pine": (495, 507, 603, 584),
        "oak": (603, 507, 710, 584),
        "note": (710, 507, 809, 584),
    }
    area_ha = 9 * (GRID_SIZE_M * GRID_SIZE_M) / 10_000
    detected = int(survey.suspicious_count)
    draw_text_fit(draw, boxes["sigungu"], record.sigungu_name, max_size=18, min_size=12, bold=True)
    draw_text_fit(
        draw,
        boxes["region"],
        f"격자 {record.center_grid_id}\n3×3 예찰권역",
        max_size=16,
        min_size=10,
    )
    draw_text_fit(draw, boxes["area"], f"{area_ha:.0f}", max_size=18, min_size=12)
    draw_text_fit(draw, boxes["total"], str(detected), max_size=20, min_size=13, bold=True)
    draw_text_fit(draw, boxes["pine"], str(detected), max_size=20, min_size=13)
    draw_text_fit(draw, boxes["oak"], "0", max_size=20, min_size=13)
    note = (
        f"시료 {survey.sample_count}점\n현장 확인 필요"
        if survey.sample_count > 0
        else "이상징후 미미\n정기 관찰"
    )
    draw_text_fit(draw, boxes["note"], note, max_size=14, min_size=10)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    image.save(output_path, format="PNG")


def verify_appendix_image_changed(template_path: Path, output_path: Path, media_name: str) -> None:
    original = extract_template_media(template_path, media_name)
    filled = Image.open(output_path).convert("RGB")
    if original.size != filled.size:
        raise RuntimeError(f"별지 이미지 크기가 바뀌었습니다: {media_name}")
    original_arr = np.asarray(original, dtype=np.int16)
    filled_arr = np.asarray(filled, dtype=np.int16)
    changed = np.mean(np.abs(original_arr - filled_arr))
    if changed < 0.15:
        raise RuntimeError(f"별지 이미지에 입력 내용이 충분히 그려지지 않았습니다: {output_path}")


def replace_docx_media(docx_path: Path, replacements: dict[str, Path]) -> None:
    """DOCX ZIP 안의 지정 media 파일을 새 PNG로 안전하게 교체한다."""
    temp_path = docx_path.with_suffix(".media-patched.docx")
    internal_replacements = {
        f"word/media/{name}": path for name, path in replacements.items()
    }

    with zipfile.ZipFile(docx_path, "r") as src, zipfile.ZipFile(
        temp_path, "w", zipfile.ZIP_DEFLATED
    ) as dst:
        names = set(src.namelist())
        missing = set(internal_replacements) - names
        if missing:
            raise RuntimeError(
                "DOCX에서 교체 대상 이미지를 찾지 못했습니다: " + ", ".join(sorted(missing))
            )
        for item in src.infolist():
            if item.filename in internal_replacements:
                dst.writestr(item, internal_replacements[item.filename].read_bytes())
            else:
                dst.writestr(item, src.read(item.filename))

    if not temp_path.exists() or temp_path.stat().st_size == 0:
        raise RuntimeError("별지 이미지 교체 후 DOCX 생성에 실패했습니다.")
    temp_path.replace(docx_path)


# -----------------------------------------------------------------------------
# DOCX 양식 채우기
# -----------------------------------------------------------------------------

def replace_in_paragraph(paragraph: Any, replacements: dict[str, str]) -> None:
    full = "".join(run.text for run in paragraph.runs)
    if not full:
        return
    new = full
    for old, value in replacements.items():
        new = new.replace(old, str(value))
    if new == full:
        return
    if paragraph.runs:
        paragraph.runs[0].text = new
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.text = new


def replace_everywhere(doc: Document, replacements: dict[str, str]) -> None:
    for paragraph in doc.paragraphs:
        replace_in_paragraph(paragraph, replacements)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_in_paragraph(paragraph, replacements)
    for section in doc.sections:
        for paragraph in section.header.paragraphs:
            replace_in_paragraph(paragraph, replacements)
        for paragraph in section.footer.paragraphs:
            replace_in_paragraph(paragraph, replacements)


def set_paragraph(doc: Document, prefix: str, text: str) -> None:
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith(prefix):
            if paragraph.runs:
                paragraph.runs[0].text = text
                for run in paragraph.runs[1:]:
                    run.text = ""
            else:
                paragraph.text = text
            return
    raise RuntimeError(f"템플릿에서 문단을 찾지 못했습니다: {prefix}")


def insert_paragraph_after(paragraph: Paragraph) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    return Paragraph(new_p, paragraph._parent)


def insert_map_after_prefix(doc: Document, prefix: str, map_path: Path) -> None:
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith(prefix):
            map_paragraph = insert_paragraph_after(paragraph)
            map_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = map_paragraph.add_run()
            run.add_picture(str(map_path), width=Mm(165))
            caption = insert_paragraph_after(map_paragraph)
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption.add_run("[현장 예찰 조사도 - 중심 격자 및 주변 3×3 권역]")
            return
    raise RuntimeError(f"지도 삽입 기준 문단을 찾지 못했습니다: {prefix}")


def create_docx(
    template_path: Path,
    output_path: Path,
    map_path: Path,
    plan_appendix_path: Path,
    result_appendix_path: Path,
    record: ReportRecord,
    cells: list[TerrainCell],
    metrics: dict[str, Any],
    survey: FieldSurveyData,
) -> None:
    doc = Document(template_path)
    region = f"{record.sido_name} {record.sigungu_name}"

    replacements = {
        "[작성일]": f"{record.year}. 12. {10 + (record.report_no % 18):02d}.",
        "-지역, 기간-": f"-{region}, {record.year}년-",
        "[일시]": survey.survey_datetime,
        "[날씨]": survey.weather,
        "기온 [ ]℃": f"기온 {survey.temperature_c:.1f}℃",
        "풍속 [ ]m/s": f"풍속 {survey.wind_speed_ms:.1f}m/s",
        "[소속·조]": survey.organization_team,
        "[성명]": survey.surveyors,
        "[입력]": survey.discovery_route,
        "[주소]": survey.address,
        "[정보]": survey.forest_compartment,
        "[위도]": f"{survey.latitude:.6f}",
        "[경도]": f"{survey.longitude:.6f}",
        "[점수·등급 또는 판단 내용]": survey.ai_result,
        "[현장 확인 필요 여부 및 판단 근거 입력]": survey.overall_judgment,
        "[수종]": survey.species,
        "[수량]": str(survey.total_trees),
        "[세부 분류]": survey.detail_classification,
        "수고 약 [ ]m": f"수고 약 {survey.tree_height_m:.1f}m",
        "흉고직경(DBH) [ ]cm": f"흉고직경(DBH) {survey.dbh_cm:.1f}cm",
        "[관찰 내용 및 단계]": survey.discoloration_stage,
        "[관찰 내용]": survey.vector_trace,
        "[현장 조사자 의견 입력]": survey.investigator_opinion,
        "[시료 종류·수량]": survey.sample_description,
        "[번호]": survey.qr_code,
        "[처리 결과 입력]": survey.system_link_result,
        "[일자]": survey.followup_date,
        "[검경 의뢰 기관 및 진행 상태]": survey.inspection_agency_status,
        "[현장 확인·검경 결과에 따른 조치 입력]": survey.followup_plan,
    }
    replace_everywhere(doc, replacements)

    # 동일한 [관찰 내용] placeholder가 2개라 명시적으로 다시 설정
    set_paragraph(doc, "❍ (변색 단계)", f"❍ (변색 단계) {survey.discoloration_stage}")
    set_paragraph(doc, "❍ (매개충 흔적)", f"❍ (매개충 흔적) {survey.vector_trace}")
    set_paragraph(doc, "❍ (수피·목질부 관찰)", f"❍ (수피·목질부 관찰) {survey.bark_wood_observation}")

    # 기존 양식의 GPS 문단 바로 뒤에 지도 삽입
    insert_map_after_prefix(doc, "― 세부 좌표:", map_path)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)

    # 템플릿의 7~8페이지 별지 이미지(image1, image2)를 채워진 이미지로 교체
    replace_docx_media(
        output_path,
        {
            "image1.png": plan_appendix_path,
            "image2.png": result_appendix_path,
        },
    )


# -----------------------------------------------------------------------------
# PDF 변환 및 검증
# -----------------------------------------------------------------------------

def find_libreoffice() -> str:
    for name in ("libreoffice", "soffice"):
        path = shutil.which(name)
        if path:
            return path
    raise RuntimeError("LibreOffice가 없습니다. sudo apt install -y libreoffice 로 설치하세요.")


def convert_docx_to_pdf(docx_path: Path, pdf_dir: Path) -> Path:
    libreoffice = find_libreoffice()
    pdf_dir.mkdir(parents=True, exist_ok=True)
    profile = pdf_dir / ".lo_profile"
    profile.mkdir(parents=True, exist_ok=True)
    command = [
        libreoffice,
        "--headless",
        f"-env:UserInstallation=file://{profile.resolve()}",
        "--convert-to",
        "pdf:writer_pdf_Export",
        "--outdir",
        str(pdf_dir),
        str(docx_path),
    ]
    result = subprocess.run(command, text=True, capture_output=True, timeout=180)
    if result.returncode != 0:
        raise RuntimeError(
            f"PDF 변환 실패: {docx_path.name}\nSTDOUT: {result.stdout}\nSTDERR: {result.stderr}"
        )
    pdf_path = pdf_dir / f"{docx_path.stem}.pdf"
    if not pdf_path.exists() or pdf_path.stat().st_size == 0:
        raise RuntimeError(f"PDF가 생성되지 않았습니다: {pdf_path}")
    return pdf_path


def verify_pdf_count(pdf_dir: Path, expected: int) -> None:
    pdfs = list(pdf_dir.glob("*.pdf"))
    if len(pdfs) != expected:
        raise RuntimeError(f"PDF가 {expected}개가 아니라 {len(pdfs)}개 생성되었습니다.")
    for pdf in pdfs:
        if pdf.stat().st_size < 10_000:
            raise RuntimeError(f"PDF 크기가 비정상적으로 작습니다: {pdf.name}")


# -----------------------------------------------------------------------------
# 메인
# -----------------------------------------------------------------------------

def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(description="기존 양식 기반 VWorld 현장 예찰 PDF 30건 생성")
    parser.add_argument("--terrain", type=Path, required=True)
    parser.add_argument("--infection", type=Path, required=True)
    parser.add_argument("--sigungu", type=Path, required=True)
    parser.add_argument("--template", type=Path, required=True)
    parser.add_argument("--prediction-manifest", type=Path, required=True, help="prediction_30/문서목록.csv")
    parser.add_argument("--output", type=Path, required=True)
    parser.add_argument("--zoom", type=int, default=None)
    parser.add_argument("--keep-docx", action="store_true", help="중간 DOCX도 보관")
    return parser


def main() -> int:
    args = build_parser().parse_args()
    project_root = Path(__file__).resolve().parents[1]
    load_dotenv(project_root / "rag-backend" / ".env", override=True)
    load_dotenv(override=False)

    api_key = os.getenv("VWORLD_API_KEY", "").strip()
    domain = os.getenv("VWORLD_API_DOMAIN", "").strip()
    basemap = os.getenv("VWORLD_BASEMAP", "GRAPHIC").strip() or "GRAPHIC"
    zoom = args.zoom or int(os.getenv("VWORLD_ZOOM", str(DEFAULT_ZOOM)))

    if not api_key:
        raise RuntimeError(".env에 VWORLD_API_KEY가 없습니다.")
    if not domain:
        raise RuntimeError(".env에 VWORLD_API_DOMAIN이 없습니다.")

    log(f"스크립트 버전: {SCRIPT_VERSION}")
    for path in (args.terrain, args.infection, args.sigungu, args.template, args.prediction_manifest):
        if not path.exists():
            raise FileNotFoundError(path)

    output_root = args.output.resolve()
    docx_dir = output_root / "docx"
    pdf_dir = output_root / "pdf"
    map_dir = output_root / "maps"
    appendix_dir = output_root / "appendices"
    output_root.mkdir(parents=True, exist_ok=True)
    docx_dir.mkdir(exist_ok=True)
    pdf_dir.mkdir(exist_ok=True)
    map_dir.mkdir(exist_ok=True)
    appendix_dir.mkdir(exist_ok=True)

    terrain_by_id, terrain_by_corner = COMMON.load_terrain_index(args.terrain)
    infection = COMMON.load_infection_history(args.infection)
    infection = COMMON.attach_admin_names(infection, args.sigungu)
    records, prediction_by_no = load_linked_prediction_records(args.prediction_manifest)
    infection_by_id = infection.set_index("id", drop=False)

    manifest: list[dict[str, Any]] = []
    total = len(records)
    if total != 30:
        log(f"주의: 예측보고서 목록이 30건이 아니라 {total}건입니다. 목록에 있는 건수만 생성합니다.")
    log(f"총 {total}건 현장 예찰 보고서 생성 시작")

    for record in records:
        log(f"  - ({record.report_no:02d}/{total}) {record.year}년 격자 {record.center_grid_id}")
        cells = COMMON.get_3x3_cells(record.center_grid_id, terrain_by_id, terrain_by_corner)
        metrics = COMMON.calculate_metrics(record, cells, infection_by_id)
        prediction_row = prediction_by_no[record.report_no]
        metrics = apply_prediction_manifest_metrics(metrics, prediction_row)
        center_cell = next(cell for cell in cells if cell.grid_id == record.center_grid_id)
        survey = create_field_survey_data(record, metrics, center_cell)

        base_name = sanitize_filename(
            f"{record.report_no:02d}_{record.year}_소나무재선충병_현장예찰보고서_"
            f"{record.sido_name}_{record.sigungu_name}_격자{record.center_grid_id}"
        )
        map_path = map_dir / f"{base_name}.png"
        plan_appendix_path = appendix_dir / f"{base_name}_07_유인항공예찰계획.png"
        result_appendix_path = appendix_dir / f"{base_name}_08_유인항공예찰조사결과.png"
        docx_path = docx_dir / f"{base_name}.docx"

        build_field_survey_map(
            output_path=map_path,
            record=record,
            cells=cells,
            infection=infection,
            survey=survey,
            api_key=api_key,
            domain=domain,
            zoom=zoom,
            basemap=basemap,
        )
        build_air_survey_plan_appendix(
            template_path=args.template,
            output_path=plan_appendix_path,
            record=record,
            survey=survey,
            metrics=metrics,
        )
        build_air_survey_result_appendix(
            template_path=args.template,
            output_path=result_appendix_path,
            record=record,
            survey=survey,
            metrics=metrics,
        )
        verify_appendix_image_changed(args.template, plan_appendix_path, "image1.png")
        verify_appendix_image_changed(args.template, result_appendix_path, "image2.png")

        create_docx(
            template_path=args.template,
            output_path=docx_path,
            map_path=map_path,
            plan_appendix_path=plan_appendix_path,
            result_appendix_path=result_appendix_path,
            record=record,
            cells=cells,
            metrics=metrics,
            survey=survey,
        )
        pdf_path = convert_docx_to_pdf(docx_path, pdf_dir)

        row = {
            "document_no": record.report_no,
            "field_report_id": f"RPT-FIELD-{record.year}-{record.report_no:03d}",
            "source_prediction_report_id": f"RPT-PRED-{record.year}-{record.report_no:03d}",
            "source_prediction_file": str(prediction_row["file_name"]),
            "file_name": pdf_path.name,
            "year": record.year,
            "center_grid_id": record.center_grid_id,
            "sido_name": record.sido_name,
            "sigungu_name": record.sigungu_name,
            "block_grid_ids": "|".join(map(str, metrics["block_grid_ids"])),
            "risk_score": round(metrics["risk_score"], 1),
            "risk_grade": metrics["risk_grade"],
            "priority_score": round(metrics["priority_score"], 1),
            "priority_grade": metrics["priority_grade"],
            "prediction_link_status": "MATCHED",
            "survey_datetime": survey.survey_datetime,
            "surveyors": survey.surveyors,
            "species": survey.species,
            "total_trees": survey.total_trees,
            "suspicious_count": survey.suspicious_count,
            "sample_count": survey.sample_count,
            "qr_code": survey.qr_code,
            "air_survey_plan_appendix": plan_appendix_path.name,
            "air_survey_result_appendix": result_appendix_path.name,
            "appendix_status": "FILLED",
            "data_origin": "historical_grid_and_rule_based_field_record",
        }
        manifest.append(row)

    verify_pdf_count(pdf_dir, total)

    manifest_path = output_root / "문서목록.csv"
    pd.DataFrame(manifest).to_csv(manifest_path, index=False, encoding="utf-8-sig")

    zip_path = output_root / f"현장예찰보고서_{total}건_예측보고서연계_VWorld_기존양식_PDF.zip"
    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as archive:
        for pdf in sorted(pdf_dir.glob("*.pdf")):
            archive.write(pdf, arcname=pdf.name)
        archive.write(manifest_path, arcname=manifest_path.name)

    if not args.keep_docx:
        shutil.rmtree(docx_dir, ignore_errors=True)

    log("완료")
    log(f"PDF 폴더: {pdf_dir}")
    log(f"ZIP 파일: {zip_path}")
    return 0


if __name__ == "__main__":
    try:
        raise SystemExit(main())
    except KeyboardInterrupt:
        print("사용자에 의해 중단되었습니다.", file=sys.stderr)
        raise SystemExit(130)
    except Exception as exc:
        print(f"오류: {exc}", file=sys.stderr)
        raise SystemExit(1)
